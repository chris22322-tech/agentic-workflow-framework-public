#!/usr/bin/env python3
"""Build Stage 6 - Evaluate - Template.docx

Generates a professional Word document template for Stage 6 of the
agentic workflow framework, matching the styling conventions established
by the Stage 3 and Stage 4 templates.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour constants (matching Stage 3/4) ──────────────────────────
NAVY = RGBColor(0x1B, 0x3A, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x66, 0x66, 0x66)
HEADER_FILL = "1B3A5C"
EXAMPLE_FILL = "D6E4F0"


# ── Helper functions ───────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_header_row(table, headers, font_size=Pt(9)):
    """Style the first row of a table as a header row."""
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = font_size
        set_cell_shading(cell, HEADER_FILL)


def add_example_row(table, row_idx, values, font_size=Pt(8)):
    """Add an example row with light blue shading."""
    for i, text in enumerate(values):
        cell = table.rows[row_idx].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = font_size
        run.italic = True
        set_cell_shading(cell, EXAMPLE_FILL)


def add_prefilled_row(table, row_idx, values, font_size=Pt(8)):
    """Add a row with pre-filled reference data (not italic, light blue shading)."""
    for i, text in enumerate(values):
        cell = table.rows[row_idx].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = font_size
        set_cell_shading(cell, EXAMPLE_FILL)


def add_empty_rows(table, count, col_count, font_size=Pt(8)):
    """Add empty rows for the user to fill in."""
    for _ in range(count):
        row = table.add_row()
        for i in range(col_count):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run("")
            run.font.size = font_size


def add_tip(doc, text):
    """Add an indented TIP callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Emu(360045)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run("TIP: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)


def add_warning(doc, text):
    """Add an indented WARNING callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Emu(360045)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run("WARNING: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)


def add_italic_instruction(doc, text):
    """Add an italic instructional paragraph the user replaces."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    return p


def add_body(doc, text, size=None):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    if size:
        for run in p.runs:
            run.font.size = size
    return p


def add_bullet(doc, text):
    """Add a bullet list item."""
    return doc.add_paragraph(text, style="List Bullet")


def add_checkbox_item(doc, text):
    """Add a checkbox item using Unicode ballot box."""
    p = doc.add_paragraph()
    run = p.add_run("\u2610  " + text)
    run.font.size = Pt(10)
    return p


def set_table_width(table, width_inches=6.5):
    """Set table alignment."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


# ── Build the document ─────────────────────────────────────────────

doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph("")

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Stage 6: Evaluate")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = NAVY

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Evaluation Report")
run.font.size = Pt(16)
run.font.color.rgb = NAVY

doc.add_paragraph("")

# Workflow Name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Workflow Name: ")
run.font.size = Pt(12)
run = p.add_run("________________________________________")
run.font.size = Pt(12)
run.font.color.rgb = GREY

doc.add_paragraph("")

# Date / Author / Version
for label in ["Date:  ", "Author:  ", "Version:  "]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(11)
    run = p.add_run("________________________")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. DOCUMENT PURPOSE
# ════════════════════════════════════════════════════════════════════

doc.add_heading("1. Document Purpose", level=1)

add_body(doc,
    "This document captures the results of evaluating your AI agent against "
    "real scenarios and comparing its output to what the manual process produces. "
    "Evaluation is not a pass/fail gate \u2014 it is a structured feedback loop. "
    "You run test cases, diagnose failures, fix the root causes, and re-test "
    "until the agent meets a quality bar you are confident in. This document "
    "records that entire journey: what you tested, what you found, what you "
    "changed, and whether the agent is ready for production use."
)

add_body(doc,
    "The evaluation report serves three purposes. First, it provides evidence "
    "that the agent works correctly across a representative range of scenarios "
    "\u2014 not just the happy path. Second, it documents the iteration cycles "
    "that improved the agent, including which failures traced back to prompts, "
    "tools, graph structure, or scope decisions. Third, it records a graduation "
    "assessment: a structured judgement, backed by evidence, about whether the "
    "agent is ready to move from prototype to production tool."
)

add_body(doc,
    "This document feeds into production deployment decisions. A completed "
    "evaluation report with all graduation criteria met is the evidence that "
    "the agent is ready to use on real work. If graduation criteria are not "
    "met, the iteration log and open questions section capture what remains "
    "to be done."
)

# ════════════════════════════════════════════════════════════════════
# 2. SCOPE
# ════════════════════════════════════════════════════════════════════

doc.add_heading("2. Scope", level=1)

doc.add_heading("What This Document Covers", level=2)

add_bullet(doc, "Test case results for a minimum of 5 representative scenarios (happy path, at-risk, data gap, sparse data, complex)")
add_bullet(doc, "Evaluation across 7 quality dimensions: correctness, completeness, quality, robustness, efficiency, HIL effectiveness, and trajectory")
add_bullet(doc, "Iteration log documenting each diagnose \u2192 categorise \u2192 fix \u2192 retest cycle")
add_bullet(doc, "Efficiency metrics comparing agent-assisted time to the manual process baseline")
add_bullet(doc, "Graduation assessment against 5 production-readiness criteria")

doc.add_heading("What This Document Does NOT Cover", level=2)

add_bullet(doc, "Production deployment configuration (infrastructure, monitoring, alerting)")
add_bullet(doc, "Ongoing production evaluation and continuous improvement after graduation")
add_bullet(doc, "User training or change management for the agent-assisted workflow")
add_bullet(doc, "Automated evaluation pipelines (e.g., LangSmith dataset runs) \u2014 this document covers manual and semi-automated evaluation")

# ════════════════════════════════════════════════════════════════════
# 3. OBJECTIVES
# ════════════════════════════════════════════════════════════════════

doc.add_heading("3. Objectives", level=1)

add_body(doc, "By completing this document, you will have:")

objectives = [
    "Validated the agent against at least 5 test cases covering different input profiles and edge cases.",
    "Assessed the agent across all 7 evaluation dimensions (correctness, completeness, quality, robustness, efficiency, HIL effectiveness, trajectory).",
    "Diagnosed and categorised every failure to its root cause layer (tool, prompt, graph, or scope).",
    "Completed at least one full iteration cycle (diagnose \u2192 categorise \u2192 fix \u2192 retest) for each identified failure.",
    "Measured efficiency gains comparing agent-assisted time to manual process time.",
    "Made an evidence-based graduation decision about whether the agent is ready for production use.",
]
for i, obj in enumerate(objectives, 1):
    add_body(doc, f"{i}. {obj}")

# ════════════════════════════════════════════════════════════════════
# 4. INPUTS REQUIRED
# ════════════════════════════════════════════════════════════════════

doc.add_heading("4. Inputs Required", level=1)

add_body(doc, "Before starting this document, you need:")

add_bullet(doc, "Working agent from Stage 5 (Build) \u2014 the agent must be runnable against test data")
add_bullet(doc, "Historical examples of the manual workflow output for comparison (at least 5 accounts/cases from the most recent period)")
add_bullet(doc, "Access to data sources the agent uses (APIs, databases, tools) with test data available")
add_bullet(doc, "The Workflow Scope Document (Stage 3) and Design Document (Stage 4) \u2014 you will reference these when diagnosing failures")
add_bullet(doc, "Time measurements for the manual process (how long each phase takes without the agent)")
add_bullet(doc, "A colleague available for blind quality comparison (optional but strongly recommended for the Quality dimension)")

# ════════════════════════════════════════════════════════════════════
# 5. OUTPUTS PRODUCED
# ════════════════════════════════════════════════════════════════════

doc.add_heading("5. Outputs Produced", level=1)

add_body(doc, "This document produces:")

add_bullet(doc, "Test case results with verdicts (Pass / Partial / Fail) and diagnostic notes for every scenario tested")
add_bullet(doc, "Evaluation dimension assessments with baseline and post-iteration scores, showing improvement over time")
add_bullet(doc, "An iteration log documenting every change made, the failure it addressed, and the retest results")
add_bullet(doc, "Efficiency metrics comparing agent-assisted time to manual process time across each workflow phase")
add_bullet(doc, "A graduation assessment \u2014 a go/no-go decision backed by evidence from the evaluation")
add_bullet(doc, "An iteration backlog of remaining improvements if the agent does not yet meet graduation criteria")

# ════════════════════════════════════════════════════════════════════
# 6. HOW TO USE THIS DOCUMENT
# ════════════════════════════════════════════════════════════════════

doc.add_heading("6. How to Use This Document", level=1)

add_body(doc,
    "Work through this document in order. Each section builds on the previous one:"
)

add_body(doc,
    "1. Start with Evaluation Approach (Section 8) \u2014 define which dimensions "
    "you are assessing and what your comparison baseline is."
)
add_body(doc,
    "2. Run your test cases and record results in Test Case Results (Section 9). "
    "Fill in one subsection per test case with the narrative and verdict."
)
add_body(doc,
    "3. For each failure, work through a full iteration cycle and record it in "
    "Iteration Log Narratives (Section 10). Diagnose, categorise, fix, retest."
)
add_body(doc,
    "4. After iterations stabilise, measure timing data and record it in "
    "Efficiency Assessment (Section 11)."
)
add_body(doc,
    "5. Complete the Graduation Assessment (Section 12) to make your "
    "production-readiness decision."
)
add_body(doc,
    "6. Review the Completion Checklist (Section 13) to verify nothing was missed."
)

add_tip(doc,
    "Tables in this document align with the companion CSV templates "
    "(stage6-test-cases.csv, stage6-evaluation-dimensions.csv, etc.). "
    "The Word document distributes some CSV columns into narrative subsections "
    "to provide richer context. Use the CSVs for structured data entry if you "
    "prefer spreadsheets, and use this document for the narrative context that "
    "surrounds the data."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 7. DEFINITIONS
# ════════════════════════════════════════════════════════════════════

doc.add_heading("7. Key Definitions", level=1)

# Failure categories table
table = doc.add_table(rows=5, cols=3, style="Table Grid")
set_table_width(table)
add_header_row(table, ["Term", "Definition", "When It Applies"])
add_example_row(table, 1, [
    "Tool Issue",
    "Data retrieval is wrong, incomplete, or failing. The agent's tools returned bad data.",
    "The raw data from an API or database is incorrect or missing. Fix the tool code."
], Pt(8))
add_example_row(table, 2, [
    "Prompt Issue",
    "The LLM is misinterpreting data or producing poorly structured output because the prompt instructions are wrong or incomplete.",
    "The data was correct but the LLM did the wrong thing with it. Fix the prompt."
], Pt(8))
add_example_row(table, 3, [
    "Graph Issue",
    "The workflow structure is wrong \u2014 nodes in the wrong order, missing conditional edges, HIL at the wrong point.",
    "Data and analysis are correct, but the workflow routes them incorrectly. Fix the graph (Stage 4)."
], Pt(8))
add_example_row(table, 4, [
    "Scope Issue",
    "The automation boundary was drawn wrong, or the scope cannot represent the distinctions the test case requires.",
    "Everything works for what it was designed to do, but the design cannot handle this scenario. Fix the scope (Stage 3)."
], Pt(8))

doc.add_paragraph("")

# Verdict definitions
table = doc.add_table(rows=4, cols=2, style="Table Grid")
set_table_width(table)
add_header_row(table, ["Verdict", "Definition"])
add_example_row(table, 1, [
    "Pass",
    "The agent produced correct output with no significant human corrections needed. Minor wording tweaks are acceptable."
], Pt(9))
add_example_row(table, 2, [
    "Partial",
    "The overall conclusion was correct but the reasoning path had a material gap, or one dimension was wrong while others were right."
], Pt(9))
add_example_row(table, 3, [
    "Fail",
    "The agent produced incorrect output that would lead to wrong actions, or a structural limitation prevented it from representing the right answer."
], Pt(9))

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 8. EVALUATION APPROACH
# ════════════════════════════════════════════════════════════════════

doc.add_heading("8. Evaluation Approach", level=1)

add_body(doc,
    "Before running test cases, define your evaluation approach: which dimensions "
    "you are assessing, what your comparison baseline is, and what \u201cgood enough\u201d "
    "looks like for each dimension."
)

doc.add_heading("8.1 Comparison Baseline", level=2)

add_italic_instruction(doc,
    "[Describe your comparison baseline. What manually-produced output will you "
    "compare the agent\u2019s output against? How many historical examples do you have? "
    "From what time period? Example: \u201cWe will compare against the manually written "
    "quarterly account health reports from Q4 2025 for 5 accounts, produced by "
    "the senior CSM who currently owns the process.\u201d]"
)

doc.add_heading("8.2 Evaluation Focus Questions", level=2)

add_italic_instruction(doc,
    "[List 2\u20134 specific questions your evaluation is designed to answer. These "
    "should be concrete and tailored to your workflow. Example: \u201c(1) Does the "
    "agent retrieve the right data from all sources? (2) Does the agent produce "
    "accurate analysis that matches manual judgement? (3) Is the final report at "
    "professional standard?\u201d]"
)

doc.add_heading("8.3 Evaluation Dimensions", level=2)

add_body(doc,
    "Assess your agent across all seven dimensions below. Skipping any of them "
    "leaves blind spots that will surface in production. The \u2018What to Assess\u2019 and "
    "\u2018How to Measure\u2019 columns are pre-filled with the framework definitions \u2014 "
    "these are reference definitions, not content you need to create. Fill in the "
    "Baseline Result after your first round of testing, and the After Iteration "
    "Result after your final iteration cycle."
)

# Evaluation Dimensions table (matches stage6-evaluation-dimensions.csv)
dim_headers = [
    "Dimension", "What to Assess", "How to Measure",
    "Baseline Result", "After Iteration Result", "Target", "Status", "Notes"
]
table = doc.add_table(rows=9, cols=8, style="Table Grid")
set_table_width(table)
add_header_row(table, dim_headers)

# All 7 dimensions pre-filled with methodology definitions
# Row 1: Correctness (example row with full worked example data)
add_example_row(table, 1, [
    "Correctness",
    "Does the agent retrieve the right data and produce accurate analysis? Are health scores, risk assessments, and trend analyses correct when compared to manual judgement?",
    "Compare agent output against manually gathered data and manually produced analysis for the same account/period. Check: same data points surfaced, same health score reached, same risks identified, same trends detected.",
    "3/5 test cases fully correct. TC2: sentiment underrated. TC5: blended score masked per-product divergence.",
    "5/5 test cases fully correct after prompt fix (TC2) and scope fix (TC5).",
    "All test cases produce correct outputs matching manual analysis",
    "Met",
    "Improved from 60% to 100% across two iteration cycles."
])

# Row 2: Completeness (pre-filled reference + empty user columns)
add_prefilled_row(table, 2, [
    "Completeness",
    "Does the agent cover all aspects the manual process covers? Are all sections of the report present? Are all data dimensions analysed?",
    "Checklist comparison against a known-good manual report. Verify: all report sections present, all data sources consulted, all risk dimensions assessed.",
    "",
    "",
    "",
    "",
    ""
])

# Row 3: Quality
add_prefilled_row(table, 3, [
    "Quality",
    "Is the written output (report narrative, executive summary, recommendations) at professional standard? Would you send it to a stakeholder after a quick review rather than a rewrite?",
    "Side-by-side blind comparison with a manually written report \u2014 have a colleague rate both without knowing which is which. Assess: narrative clarity, appropriate tone, actionable recommendations.",
    "",
    "",
    "",
    "",
    ""
])

# Row 4: Robustness
add_prefilled_row(table, 4, [
    "Robustness",
    "How does the agent handle missing data, API failures, unusual inputs, and edge cases? Does it degrade gracefully or crash/hallucinate?",
    "Run against edge case accounts: API failure, sparse data/new account, contradictory signals. Verify: no crashes, no hallucinated data, explicit flagging of gaps, appropriate caveats on low-confidence assessments.",
    "",
    "",
    "",
    "",
    ""
])

# Row 5: Efficiency
add_prefilled_row(table, 5, [
    "Efficiency",
    "Is the agent actually saving time when you account for both agent execution time and the human time spent reviewing and editing the output at each checkpoint?",
    "Measure total time (agent execution time + human review/edit time at each HIL checkpoint) vs. manual process time for the same account. Target: total agent-assisted time < 50% of manual process time.",
    "",
    "",
    "",
    "",
    ""
])

# Row 6: HIL Effectiveness
add_prefilled_row(table, 6, [
    "HIL Effectiveness",
    "Are the human-in-the-loop checkpoints at the right places? Is the information surfaced at each checkpoint useful for the reviewer? Are checkpoints adding value or just adding delay?",
    "Track the size of human edits at each checkpoint across all test cases. Categories: none, minor tweaks, significant corrections, complete rewrite. If edits are consistently none/minor, the checkpoint may be unnecessary. If consistently significant/rewrite, the preceding node needs improvement.",
    "",
    "",
    "",
    "",
    ""
])

# Row 7: Trajectory
add_prefilled_row(table, 7, [
    "Trajectory",
    "Did the agent take the correct path through the graph \u2014 right node sequence, no unnecessary loops, no skipped steps? Did it reach the right output via the right process, not just by coincidence?",
    "Define expected node sequences for each test case and compare against actual execution order via traces (LangSmith or manual logging). Flag cases where the agent reached the correct output via an inefficient or incorrect path.",
    "",
    "",
    "",
    "",
    ""
])

# One blank user row for a custom dimension
add_empty_rows(table, 1, 8)

add_tip(doc,
    "Trajectory evaluation catches failures invisible to output-only assessment. "
    "An agent can produce a correct report while taking an unnecessarily expensive "
    "path \u2014 re-running analysis after an unnecessary loop, skipping a data source "
    "and compensating downstream. Define expected node sequences for each test case "
    "and compare against actual execution traces."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 9. TEST CASE RESULTS
# ════════════════════════════════════════════════════════════════════

doc.add_heading("9. Test Case Results", level=1)

add_body(doc,
    "Create a minimum of 5 test cases representing different input profiles. "
    "Do not just test the happy path \u2014 the edge cases are where agents fail. "
    "For each test case, record the input, expected behaviour, actual behaviour, "
    "verdict, and diagnostic notes."
)

add_body(doc,
    "The table below provides a summary view. Each test case also has a detailed "
    "narrative subsection below the table where you write up the full results."
)

add_warning(doc,
    "Only testing the happy path is the most common evaluation mistake. The agent "
    "will look great on clean data with no failures. Test the edge cases \u2014 sparse "
    "data, API errors, ambiguous inputs \u2014 because those are the conditions it will "
    "encounter in production."
)

# Test Cases Summary Table
tc_headers = [
    "Test Case ID", "Test Case Name", "Account Profile",
    "Purpose", "Verdict", "Failure Category", "Notes"
]
table = doc.add_table(rows=7, cols=7, style="Table Grid")
set_table_width(table)
add_header_row(table, tc_headers)

# Example row
add_example_row(table, 1, [
    "TC1",
    "Happy Path \u2014 Healthy Account",
    "Acme Corp (ACME-001). Healthy account, all data available, no anomalies.",
    "Verify the basic end-to-end flow works. All data sources return clean data.",
    "Pass",
    "N/A",
    "Exec summary slightly generic \u2014 minor quality gap, not a correctness issue."
])

# Pre-fill test case type names in remaining rows
tc_types = [
    "TC2 \u2014 At-Risk Scenario",
    "TC3 \u2014 Data Gap",
    "TC4 \u2014 Sparse Data",
    "TC5 \u2014 Complex Case",
    "TC6 \u2014 [Additional]",
]
for idx, name in enumerate(tc_types, 2):
    cell = table.rows[idx].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"TC{idx}")
    run.font.size = Pt(8)
    run.italic = True
    cell2 = table.rows[idx].cells[1]
    cell2.text = ""
    p2 = cell2.paragraphs[0]
    run2 = p2.add_run(name.split(" \u2014 ")[1] if " \u2014 " in name else "")
    run2.font.size = Pt(8)
    run2.italic = True

add_empty_rows(table, 3, 7)

doc.add_paragraph("")

# ── Detailed Test Case Narratives ──────────────────────────────────

tc_detail_headers = ["Field", "Details"]

# ── TC1: Happy Path (worked example — Pass verdict)
doc.add_heading("9.1 Test Case 1: Happy Path", level=2)

add_body(doc, "Purpose: Verify the basic end-to-end flow works with clean data and no anomalies.")

table = doc.add_table(rows=6, cols=2, style="Table Grid")
set_table_width(table)
add_header_row(table, tc_detail_headers)
add_example_row(table, 1, [
    "Input",
    "account_id: ACME-001, Q4 data. Healthy account: stable usage (including 15% usage increase), "
    "no open P1 tickets, Green health score history. All APIs returning normally."
])
add_example_row(table, 2, [
    "Expected Behaviour",
    "Agent gathers all data cleanly. Produces a Green health score with correct justification. "
    "Flags no major risks. Generates a professional report with minor or no edits at HIL checkpoints."
])
add_example_row(table, 3, [
    "Actual Behaviour",
    "Agent gathered all data, produced Green health score with correct justification. Report "
    "quality comparable to manually written version. Zero changes at hil_review_analysis; one "
    "minor wording tweak at hil_review_report."
])
add_example_row(table, 4, [
    "Verdict & Rationale",
    "Pass. All data retrieved correctly, health score matched manual assessment (Green), "
    "risk flags appropriate (none), report at professional standard."
])
add_example_row(table, 5, [
    "Issues & Notes",
    "Exec summary slightly generic (\u201caccount remains healthy\u201d) \u2014 a manually written "
    "summary would have led with the 15% usage increase. Minor quality gap in narrative "
    "framing, not a correctness issue."
])

doc.add_paragraph("")

# ── TC2: At-Risk (worked example — Partial verdict, shows non-Pass documentation)
doc.add_heading("9.2 Test Case 2: At-Risk Scenario", level=2)

add_body(doc, "Purpose: Verify the agent correctly identifies and escalates risk signals.")

table = doc.add_table(rows=6, cols=2, style="Table Grid")
set_table_width(table)
add_header_row(table, tc_detail_headers)
add_example_row(table, 1, [
    "Input",
    "account_id: GLOBEX-042, Q4 data. Usage down 30% QoQ, 3 open P1 tickets, renewal in "
    "60 days, negative Slack sentiment including clearly frustrated messages from the VP of "
    "Engineering alongside general end-user complaints."
])
add_example_row(table, 2, [
    "Expected Behaviour",
    "Agent flags Red or Amber health score. Identifies churn risk with specific evidence "
    "(usage decline, P1 tickets, negative sentiment). Surfaces the 60-day renewal timeline "
    "as a risk amplifier. Sentiment analysis recognises executive-level frustration as higher risk."
])
add_example_row(table, 3, [
    "Actual Behaviour",
    "Agent correctly scored Red, identified churn risk, and cited usage decline and open P1s. "
    "However, sentiment analysis scored \u201cNeeds Attention\u201d instead of \u201cCritical\u201d \u2014 the "
    "SENTIMENT_ANALYSIS_PROMPT treated all messages with equal weight regardless of stakeholder "
    "seniority. The report did not distinguish executive frustration from end-user complaints."
])
add_example_row(table, 4, [
    "Verdict & Rationale",
    "Partial. Health score and churn risk identification were correct (Red, with usage and "
    "ticket evidence), but sentiment severity was underrated. For this account, other signals "
    "were strong enough to produce the correct Red score, but for a borderline account this "
    "gap would produce the wrong verdict. Partial because overall conclusion correct but "
    "reasoning path had a material gap."
])
add_example_row(table, 5, [
    "Issues & Notes",
    "Root cause: Prompt issue. SENTIMENT_ANALYSIS_PROMPT has no guidance on weighting by "
    "stakeholder role or seniority. VP of Engineering frustration near a renewal is a "
    "materially higher risk signal than end-user complaints. Fixed in Iteration 1 by adding "
    "stakeholder weighting instructions. After fix, TC2 re-tested as Pass."
])

add_tip(doc,
    "This example shows how to document a Partial verdict. Note how the rationale explains "
    "why the test case was not a full Pass (sentiment underrated) even though the overall "
    "conclusion (Red health score) was correct. The Notes section identifies the failure "
    "category (Prompt) and references the iteration that fixed it."
)

doc.add_paragraph("")

# ── TC3 template
doc.add_heading("9.3 Test Case 3: Data Gap", level=2)

add_body(doc, "Purpose: Verify graceful degradation when a data source fails or returns empty.")

table = doc.add_table(rows=6, cols=2, style="Table Grid")
set_table_width(table)
add_header_row(table, tc_detail_headers)
for row_idx in range(1, 6):
    labels = ["Input", "Expected Behaviour", "Actual Behaviour", "Verdict & Rationale", "Issues & Notes"]
    cell = table.rows[row_idx].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(labels[row_idx - 1])
    run.font.size = Pt(8)
    run.italic = True

add_italic_instruction(doc,
    "[Describe which data source you disabled or simulated an error for, how the "
    "agent handled the gap (did it flag it explicitly? hallucinate? crash?), and your "
    "verdict. The target behaviour is explicit flagging with no hallucination.]"
)

# ── TC4 template
doc.add_heading("9.4 Test Case 4: Sparse Data", level=2)

add_body(doc, "Purpose: Verify the agent does not over-interpret limited data.")

table = doc.add_table(rows=6, cols=2, style="Table Grid")
set_table_width(table)
add_header_row(table, tc_detail_headers)
for row_idx in range(1, 6):
    labels = ["Input", "Expected Behaviour", "Actual Behaviour", "Verdict & Rationale", "Issues & Notes"]
    cell = table.rows[row_idx].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(labels[row_idx - 1])
    run.font.size = Pt(8)
    run.italic = True

add_italic_instruction(doc,
    "[Describe the sparse-data scenario (e.g., new account with limited history), "
    "whether the agent qualified its confidence appropriately, and your verdict. "
    "The agent should note data limitations rather than producing a definitive "
    "assessment from insufficient data.]"
)

# ── TC5 template
doc.add_heading("9.5 Test Case 5: Complex Case", level=2)

add_body(doc, "Purpose: Verify the agent handles nuance and does not oversimplify mixed signals.")

table = doc.add_table(rows=6, cols=2, style="Table Grid")
set_table_width(table)
add_header_row(table, tc_detail_headers)
for row_idx in range(1, 6):
    labels = ["Input", "Expected Behaviour", "Actual Behaviour", "Verdict & Rationale", "Issues & Notes"]
    cell = table.rows[row_idx].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(labels[row_idx - 1])
    run.font.size = Pt(8)
    run.italic = True

add_italic_instruction(doc,
    "[Describe the complex scenario (e.g., multiple products, mixed signals across "
    "dimensions, ambiguous indicators). Did the agent surface the complexity rather "
    "than flattening it into a single score? Were recommendations specific to the "
    "nuances of this case, or generic?]"
)

# ── Additional test case slots
doc.add_heading("9.6 Additional Test Cases", level=2)

add_italic_instruction(doc,
    "[Add additional test cases as needed. Copy the table structure from any test "
    "case above. You should have a minimum of 5 test cases, but more is better \u2014 "
    "especially if your workflow handles diverse input profiles.]"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 10. ITERATION LOG NARRATIVES
# ════════════════════════════════════════════════════════════════════

doc.add_heading("10. Iteration Log Narratives", level=1)

add_body(doc,
    "After each evaluation round, follow the iteration cycle: diagnose the failure, "
    "categorise it (tool / prompt / graph / scope), implement the fix, and re-test. "
    "Record each iteration below. The summary table tracks progress at a glance; the "
    "narrative subsections capture the full diagnose \u2192 fix \u2192 retest story for each "
    "iteration."
)

add_warning(doc,
    "Iterating on prompts without tracking what you changed is a common mistake. "
    "Prompt tuning is empirical. If you tweak a prompt and re-run without recording "
    "the change and result, you will lose track of what works and end up going in circles."
)

doc.add_heading("10.1 Iteration Summary Table", level=2)

add_body(doc,
    "This table provides an at-a-glance view of each iteration. The narrative "
    "subsections below capture the full diagnostic detail, code changes, and "
    "upstream cascades that the summary table condenses."
)

# Iteration Log table — includes Files Modified and Regression Check
# for engineering traceability
iter_headers = [
    "Iteration #", "Date", "Trigger",
    "Failure Category", "Change Made", "Files Modified",
    "Test Cases Re-Run", "Results", "Regression Check", "Notes"
]
table = doc.add_table(rows=2, cols=10, style="Table Grid")
set_table_width(table)
add_header_row(table, iter_headers)

add_example_row(table, 1, [
    "1",
    "2026-03-15",
    "TC2: Partial. Sentiment scored \u201cNeeds Attention\u201d instead of \u201cCritical.\u201d",
    "Prompt",
    "Added stakeholder seniority weighting to SENTIMENT_ANALYSIS_PROMPT.",
    "prompts.py",
    "TC2, TC1, TC3, TC4",
    "TC2: Partial \u2192 Pass. TC1,3,4: no regression.",
    "TC1: Pass. TC3: Pass. TC4: Pass. No negative effect.",
    "Weighting uses available role signals; degrades gracefully when absent."
])

add_empty_rows(table, 5, 10)

doc.add_paragraph("")

# ── Iteration Narrative: Example

doc.add_heading("10.2 Iteration 1 Narrative", level=2)

add_body(doc, "Record the full story of each iteration cycle below.")

doc.add_heading("Diagnose", level=3)

add_italic_instruction(doc,
    "[What failure did you observe? What was the expected vs actual behaviour? "
    "Be specific about the gap \u2014 e.g., \u201cSentiment analysis returned \u2018Needs Attention\u2019 "
    "for an account where the VP of Engineering expressed clear frustration. The "
    "manually written report rated sentiment as \u2018Critical.\u2019\u201d]"
)

doc.add_heading("Categorise", level=3)

add_body(doc,
    "Use systematic elimination to find the root cause. Check each layer in order:"
)

add_body(doc, "1. Tool layer \u2014 Was the right data retrieved?")
add_body(doc, "2. Prompt layer \u2014 Given the data it received, did the prompt produce the right analysis?")
add_body(doc, "3. Graph layer \u2014 Is the node in the right position with the right inputs?")
add_body(doc, "4. Scope layer \u2014 Does the scope support what the agent needs to represent?")

add_italic_instruction(doc,
    "[Record your systematic elimination. Which layers passed? Which layer failed? "
    "What is the failure category? Example: \u201cTool \u2713 (correct Slack threads retrieved). "
    "Prompt \u2717 (no guidance on stakeholder weighting). This is a prompt issue.\u201d]"
)

doc.add_heading("Fix", level=3)

add_italic_instruction(doc,
    "[Describe the specific change you made. Include before/after for prompt changes. "
    "For code changes, name the file and function modified. For scope or graph changes, "
    "describe the upstream artifact updates and the downstream cascade.]"
)

doc.add_heading("Retest", level=3)

add_italic_instruction(doc,
    "[Which test cases did you re-run? What were the results? Did the fix resolve "
    "the original failure? Did it cause any regressions on other test cases?]"
)

add_tip(doc,
    "For graph or scope failures, you must cascade changes forward through each "
    "downstream stage in order. A scope change that bypasses the Design Document "
    "creates drift between what the design says and what the code does."
)

# ── Additional iteration slots
doc.add_heading("10.3 Iteration 2 Narrative", level=2)

add_italic_instruction(doc,
    "[Copy the Diagnose / Categorise / Fix / Retest structure from Iteration 1 above. "
    "If this iteration involves a cross-stage cascade (graph or scope issue), include "
    "a subsection documenting which upstream artifacts were updated and how the changes "
    "cascaded forward through Design and Build.]"
)

doc.add_heading("10.4 Additional Iterations", level=2)

add_italic_instruction(doc,
    "[Add additional iteration narratives as needed. Each follows the same structure: "
    "Diagnose, Categorise, Fix, Retest. Number them sequentially.]"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 11. EFFICIENCY ASSESSMENT
# ════════════════════════════════════════════════════════════════════

doc.add_heading("11. Efficiency Assessment", level=1)

add_body(doc,
    "Time the full agent-assisted process (agent execution + your review time at "
    "each HIL checkpoint) and compare against the manual process time. The target: "
    "total agent-assisted time should be less than 50% of the manual process time."
)

add_body(doc,
    "Fill in the table below with your measurements. The example row shows the "
    "format and level of detail expected."
)

# Efficiency Metrics table — includes Target column (matching CSV)
eff_headers = [
    "Metric", "Manual Process Time", "Agent Execution Time",
    "Human Review Time", "Total Agent-Assisted",
    "Time Saved", "% Improvement", "Target", "Status"
]
table = doc.add_table(rows=7, cols=9, style="Table Grid")
set_table_width(table)
add_header_row(table, eff_headers)

# Example row
add_example_row(table, 1, [
    "Total time per account",
    "7 hrs (avg 6\u20138 hr range)",
    "12 min",
    "1 hr 30 min",
    "1 hr 42 min",
    "5 hrs 18 min",
    "76%",
    "< 3 hrs 30 min (50% of manual)",
    "Met"
])

# Pre-fill metric labels with targets
metric_rows = [
    ("Data gathering", "< 15 min human review"),
    ("Analysis", "< 30 min human review"),
    ("Report writing", "< 30 min human editing"),
    ("HIL Checkpoint 1", "Edits none/minor in > 80% of cases"),
    ("HIL Checkpoint 2", "Edits none/minor in > 80% of cases"),
]
for idx, (label, target) in enumerate(metric_rows, 2):
    cell = table.rows[idx].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.font.size = Pt(8)
    run.italic = True
    cell_target = table.rows[idx].cells[7]
    cell_target.text = ""
    p_t = cell_target.paragraphs[0]
    run_t = p_t.add_run(target)
    run_t.font.size = Pt(8)
    run_t.italic = True

add_empty_rows(table, 3, 9)

doc.add_paragraph("")

doc.add_heading("11.1 Efficiency Narrative", level=2)

add_italic_instruction(doc,
    "[Interpret your timing data. Where are the biggest time savings? Where is "
    "human review time still high (indicating prompt quality improvements could "
    "help)? Is the 50% target met? If not, what would need to improve to meet "
    "it? Example: \u201cThe largest time saving is in data gathering (from 2.5 hrs "
    "to 18 min). Human review time is the dominant component of the agent-assisted "
    "total \u2014 further prompt improvements could reduce review time at checkpoints.\u201d]"
)

add_tip(doc,
    "If total agent-assisted time exceeds 50% of the manual process, focus on the "
    "phase where human review time is highest. That is where prompt quality "
    "improvements will have the most impact on overall efficiency."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 12. GRADUATION ASSESSMENT
# ════════════════════════════════════════════════════════════════════

doc.add_heading("12. Graduation Assessment", level=1)

add_body(doc,
    "The agent is ready for production use when all five graduation criteria are "
    "met. Assess each criterion below with evidence from your evaluation."
)

# Graduation Checklist table (matches stage6-graduation-checklist.csv — 7 columns)
grad_headers = [
    "Criterion", "Description", "Status",
    "Evidence", "Date Verified", "Verified By", "Notes"
]
table = doc.add_table(rows=6, cols=7, style="Table Grid")
set_table_width(table)
add_header_row(table, grad_headers)

# Example row
add_example_row(table, 1, [
    "All test cases pass without major corrections",
    "All 5+ test cases pass with no significant corrections or complete rewrites at HIL checkpoints.",
    "Pass",
    "5/5 test cases pass after two iteration cycles. HIL edits now all none/minor tweaks.",
    "2026-03-17",
    "[Name]",
    ""
])

# Pre-fill criterion names and descriptions
criteria = [
    ("Total time under 50% of manual",
     "Agent-assisted time (execution + review) < 50% of manual process time."),
    ("No critical failures in last 3 rounds",
     "No wrong conclusions or missed major risks in the last 3 consecutive evaluation rounds."),
    ("Colleague review of output quality",
     "At least one colleague has reviewed and validated output quality via blind comparison."),
    ("Error handling tested for all failure modes",
     "Error handling tested for all identified failure modes (API failures, missing data, sparse data, contradictory signals)."),
]
for idx, (name, desc) in enumerate(criteria, 2):
    cell_name = table.rows[idx].cells[0]
    cell_name.text = ""
    p = cell_name.paragraphs[0]
    run = p.add_run(name)
    run.font.size = Pt(8)
    run.italic = True
    cell_desc = table.rows[idx].cells[1]
    cell_desc.text = ""
    p2 = cell_desc.paragraphs[0]
    run2 = p2.add_run(desc)
    run2.font.size = Pt(8)
    run2.italic = True

doc.add_paragraph("")

doc.add_heading("12.1 Overall Graduation Decision", level=2)

add_italic_instruction(doc,
    "[State your graduation decision: does the agent graduate to production use, "
    "or does it require further iteration? Summarise the evidence. If the agent "
    "does not graduate, list the specific criteria that are not yet met and what "
    "needs to happen to meet them.]"
)

add_warning(doc,
    "Moving too fast to remove HIL checkpoints is a common mistake. It is tempting "
    "to remove human review once the agent seems reliable. Wait for 2\u20133 clean "
    "evaluation cycles first. Early success can be coincidence; sustained success "
    "is evidence."
)

doc.add_heading("12.2 Automation Boundary Recommendations", level=2)

add_body(doc,
    "After 2\u20133 successful evaluation cycles with consistent results, consider "
    "whether the automation boundary should move:"
)

add_italic_instruction(doc,
    "[Based on your HIL checkpoint edit patterns, are there checkpoints that "
    "could be removed (human rarely makes changes)? Are there MANUAL steps that "
    "turned out to be more formulaic than expected and could become HIL steps? "
    "Are there new data sources that could be added? Document any recommendations "
    "here with supporting evidence from your evaluation data.]"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 13. COMPLETION CHECKLIST
# ════════════════════════════════════════════════════════════════════

doc.add_heading("13. Completion Checklist", level=1)

add_body(doc,
    "Before considering this document complete, verify that you have addressed "
    "every item below."
)

checklist_items = [
    "Comparison baseline defined (Section 8.1) \u2014 manual output identified for comparison",
    "Evaluation dimensions assessed (Section 8.3) \u2014 all 7 dimensions have baseline and post-iteration results",
    "At least 5 test cases run (Section 9) \u2014 including happy path, at-risk, data gap, sparse data, and complex scenarios",
    "Every test case has a detailed narrative (Section 9) \u2014 input, expected, actual, verdict, and notes filled in",
    "Every Partial or Fail verdict has a diagnosed root cause with failure category identified",
    "At least one full iteration cycle completed (Section 10) \u2014 diagnose, categorise, fix, retest",
    "Iteration log records every change made, with before/after details and retest results",
    "Regression checks performed \u2014 every iteration re-ran unaffected test cases to confirm no regression",
    "Efficiency metrics recorded (Section 11) \u2014 manual vs agent-assisted time for each phase",
    "Graduation criteria assessed (Section 12) \u2014 all 5 criteria have status, evidence, and date",
    "Overall graduation decision stated (Section 12.1) with supporting evidence",
    "Open questions captured (Section 14) \u2014 anything unresolved or uncertain is documented",
]

for item in checklist_items:
    add_checkbox_item(doc, item)

# ════════════════════════════════════════════════════════════════════
# 14. NOTES & OPEN QUESTIONS
# ════════════════════════════════════════════════════════════════════

doc.add_heading("14. Notes & Open Questions", level=1)

add_italic_instruction(doc,
    "[Use this section to capture anything you are unsure about, edge cases "
    "you haven\u2019t tested yet, potential improvements for future iterations, "
    "or questions that arose during evaluation. This section ensures nothing "
    "falls through the cracks.]"
)

add_body(doc, "")
add_body(doc, "")
add_body(doc, "")

# ════════════════════════════════════════════════════════════════════
# 15. NEXT STEPS
# ════════════════════════════════════════════════════════════════════

doc.add_heading("15. Next Steps", level=1)

add_body(doc,
    "Once this document is complete:"
)

add_body(doc,
    "If the agent graduated: The agent is ready for production use. Begin using "
    "it on real work, starting with a subset of cases while you build confidence. "
    "Continue tracking HIL checkpoint edit patterns as production evaluation data "
    "\u2014 approvals, modifications, and rejections all provide signal for future "
    "improvement. Over time, use this data to decide when to expand the automation "
    "boundary (remove a checkpoint, add a new data source, or bring a MANUAL step "
    "into the graph)."
)

add_body(doc,
    "If the agent did not graduate: Review the open items in the graduation "
    "assessment (Section 12) and the Notes & Open Questions (Section 14). Prioritise "
    "remaining fixes using Impact \u00d7 Feasibility. Continue the iteration cycle "
    "(Section 10) until all graduation criteria are met, then re-assess."
)

add_body(doc,
    "In either case, store this evaluation report alongside the Workflow Scope "
    "Document (Stage 3), Design Document (Stage 4), and agent code (Stage 5). "
    "Together, these four artifacts form a complete record of what the agent does, "
    "why it was designed that way, how it was built, and the evidence that it works."
)

add_tip(doc,
    "This is the final stage of the framework methodology. For reference material "
    "on LangGraph patterns, decision trees, and checklists, see the Reference "
    "section of the documentation."
)

# ════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════

output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, "Stage 6 - Evaluate - Template.docx")
doc.save(output_path)
print(f"Created: {output_path}")
