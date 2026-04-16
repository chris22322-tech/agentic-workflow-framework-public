#!/usr/bin/env python3
"""Generate Stage 3 - Scope - Template.docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour constants ──────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
LIGHT_BLUE_HEX = "D6E4F0"
DARK_BLUE_HEX = "1B3A5C"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TIP_GREEN = RGBColor(0x1E, 0x7D, 0x32)
WARNING_RED = RGBColor(0xC6, 0x28, 0x28)
GREY = RGBColor(0x66, 0x66, 0x66)


def set_cell_shading(cell, colour_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colour_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, bg_hex=DARK_BLUE_HEX):
    """Style a table header row: dark background, white bold text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_hex)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.bold = True
                run.font.size = Pt(9)


def style_example_row(row, bg_hex=LIGHT_BLUE_HEX):
    """Style an example row with light blue background."""
    for cell in row.cells:
        set_cell_shading(cell, bg_hex)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8.5)


def add_table_with_headers(doc, headers, col_widths=None):
    """Add a table with styled header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    style_header_row(table.rows[0])

    # Apply column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_example_row(table, values):
    """Add a pre-filled example row with light blue shading."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(8.5)
    style_example_row(row)
    return row


def add_blank_row(table, num_cols, placeholder=""):
    """Add a blank row for user input."""
    row = table.add_row()
    for i in range(num_cols):
        row.cells[i].text = placeholder
    return row


def add_italic_instruction(doc, text):
    """Add instructional text in italics."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = GREY
    run.font.size = Pt(10)
    return p


def add_tip(doc, text):
    """Add a TIP callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run_label = p.add_run("TIP: ")
    run_label.bold = True
    run_label.font.color.rgb = TIP_GREEN
    run_label.font.size = Pt(10)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)
    run_body.font.color.rgb = TIP_GREEN
    return p


def add_warning(doc, text):
    """Add a WARNING callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run_label = p.add_run("WARNING: ")
    run_label.bold = True
    run_label.font.color.rgb = WARNING_RED
    run_label.font.size = Pt(10)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)
    run_body.font.color.rgb = WARNING_RED
    return p


def add_checkbox_item(doc, text):
    """Add a checklist item with a checkbox character."""
    p = doc.add_paragraph()
    run = p.add_run("\u2610  " + text)
    run.font.size = Pt(10)
    return p


# ══════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════

doc = Document()

# ── Default font ──────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10.5)

for heading_level in range(1, 4):
    hs = doc.styles[f"Heading {heading_level}"]
    hs.font.color.rgb = DARK_BLUE

# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Stage 3: Scope")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Workflow Scope Document")
run.font.size = Pt(16)
run.font.color.rgb = DARK_BLUE

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Workflow Name: ")
run.font.size = Pt(12)
run2 = p.add_run("________________________________________")
run2.font.size = Pt(12)
run2.italic = True
run2.font.color.rgb = GREY

doc.add_paragraph("")

for label in ["Date:", "Author:", "Version:"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}  ")
    run.font.size = Pt(11)
    run2 = p.add_run("________________________")
    run2.font.size = Pt(11)
    run2.italic = True
    run2.font.color.rgb = GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. DOCUMENT PURPOSE
# ══════════════════════════════════════════════════════════════════

doc.add_heading("1. Document Purpose", level=1)

doc.add_paragraph(
    "This document captures the detailed scope of a single workflow selected for AI agent automation. "
    "It translates your domain expertise — the steps you take, the judgement calls you make, the data "
    "you consult — into a structured format that can be handed off to the Design and Build stages of "
    "the framework."
)

doc.add_paragraph(
    "The Scope Document is the single most important artifact in the framework. A well-scoped workflow "
    "is straightforward to design and build. A poorly scoped one produces an agent that either does too "
    "little to be useful or too much to be trustworthy. This is where your domain expertise matters "
    "most: you are the person who knows the real workflow — the shortcuts, the judgement calls, the "
    "things nobody writes down."
)

doc.add_paragraph(
    "Once completed, this document feeds directly into Stage 4 (Design), where every step, data source, "
    "and decision point you record here is translated into a LangGraph agent architecture. The quality "
    "of the design is entirely determined by the quality of this scope."
)

# ══════════════════════════════════════════════════════════════════
# 2. SCOPE
# ══════════════════════════════════════════════════════════════════

doc.add_heading("2. Scope", level=1)

doc.add_heading("What This Document Covers", level=2)

doc.add_paragraph(
    "This document covers the complete operational scope of a single workflow selected in Stage 2, including:"
)

items_in = [
    "Every step a human currently takes to execute the workflow, in sequence",
    "The data sources consulted at each step, with access methods and formats",
    "The decision logic applied at each step — thresholds, criteria, and contextual judgement",
    "The automation boundary for each step (AUTOMATE, HUMAN-IN-THE-LOOP, or MANUAL)",
    "Error and edge case paths for each step",
    "Integration requirements (APIs, libraries, credentials, infrastructure)",
    "Constraints and assumptions that must hold for the workflow to operate as designed",
]
for item in items_in:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("What This Document Does NOT Cover", level=2)

items_out = [
    "Agent architecture or graph design (Stage 4: Design)",
    "Implementation details, code, or prompt engineering (Stage 5: Build)",
    "Workflow selection rationale (Stage 2: Select — assumed complete)",
    "Batch processing or multi-workflow orchestration",
    "Deployment, monitoring, or operational runbooks (Stage 6: Stabilise)",
]
for item in items_out:
    doc.add_paragraph(item, style="List Bullet")

# ══════════════════════════════════════════════════════════════════
# 3. OBJECTIVES
# ══════════════════════════════════════════════════════════════════

doc.add_heading("3. Objectives", level=1)

doc.add_paragraph(
    "By completing this document, you will have accomplished the following:"
)

objectives = [
    "Mapped every step of the workflow at a level of detail sufficient for a new hire — or an AI agent — to follow.",
    "Captured the decision logic at each step, distinguishing between criteria-based decisions (translatable to config) and contextual judgement (translatable to prompt instructions).",
    "Drawn the automation boundary: each step tagged as AUTOMATE, HUMAN-IN-THE-LOOP, or MANUAL, with a rationale for each assignment.",
    "Documented error and edge case paths so failures are handled by design, not discovered during build.",
    "Inventoried every data source the workflow consumes, with access method, format, and authentication requirements.",
    "Listed all integration requirements (frameworks, APIs, credentials, output tooling) needed to build the agent.",
    "Surfaced hidden assumptions using the five diagnostic categories and stated the specific values you are assuming.",
]
for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    run2 = p.add_run(obj)

# ══════════════════════════════════════════════════════════════════
# 4. INPUTS REQUIRED
# ══════════════════════════════════════════════════════════════════

doc.add_heading("4. Inputs Required", level=1)

doc.add_paragraph(
    "Before starting this document, ensure you have the following:"
)

inputs = [
    ("Selection Decision Record from Stage 2", "The completed Stage 2 artifact confirming which workflow was selected and why."),
    ("Your own knowledge of the workflow", "You must be someone who currently executes this workflow (or has done so recently). The value of this document comes from capturing tacit knowledge — the things you do on autopilot that are not written down anywhere."),
    ("Access to the systems involved", "You will need to reference the actual tools, dashboards, and data sources you use when performing the workflow, to accurately document access methods and data formats."),
    ("SLA targets, scoring frameworks, or business rules (if applicable)", "If your workflow involves evaluating data against defined criteria (e.g., health scores, priority tiers, compliance thresholds), have those definitions available."),
]
for title, desc in inputs:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════════════════════════════
# 5. OUTPUTS PRODUCED
# ══════════════════════════════════════════════════════════════════

doc.add_heading("5. Outputs Produced", level=1)

doc.add_paragraph("This document produces the following artifacts that Stage 4 (Design) consumes:")

outputs = [
    ("Step-by-Step Workflow Map", "The sequence of steps, with decision logic and boundary tags, that becomes the node/edge structure of the LangGraph agent."),
    ("Data Inventory", "The list of data sources that becomes the tool specification for each data-retrieval node."),
    ("Integration Requirements", "The technology stack that informs the build environment and dependency setup."),
    ("Constraints and Assumptions", "The operating conditions that shape the state schema, error handling strategy, and scope limitations of the agent design."),
]
for title, desc in outputs:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════════════════════════════
# 6. HOW TO USE THIS DOCUMENT
# ══════════════════════════════════════════════════════════════════

doc.add_heading("6. How to Use This Document", level=1)

instructions = [
    "Read the guidance text in each section before filling anything in. The guidance explains what level of detail is expected and why.",
    "Use the example rows (shaded in light blue) as a reference for the depth and specificity required. Your entries should match or exceed this level of detail.",
    "Fill in the tables first, then write the narrative sections. The tables capture structured data; the narratives capture context, reasoning, and edge cases that tables cannot hold.",
    "Replace all italic instructional text with your own content. If a section does not apply to your workflow, write a brief note explaining why rather than leaving it blank.",
    "Review your completed document against the Completion Checklist at the end. Every item should be ticked before you hand this document to the person responsible for Stage 4.",
]
for i, instr in enumerate(instructions, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(instr)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7. WORKFLOW OVERVIEW
# ══════════════════════════════════════════════════════════════════

doc.add_heading("7. Workflow Overview", level=1)

# Workflow name
p = doc.add_paragraph()
run = p.add_run("Workflow Name: ")
run.bold = True
add_italic_instruction(doc, "Enter the name of the workflow exactly as it appears in your Stage 2 Selection Decision Record.")

# Source
p = doc.add_paragraph()
run = p.add_run("Source (Stage 2 Reference): ")
run.bold = True
add_italic_instruction(doc, "Reference the Stage 2 document — e.g., 'Selected as highest-scoring workflow in Stage 2 Selection Matrix, dated [date].'")

# Scope summary
p = doc.add_paragraph()
run = p.add_run("Scope Summary: ")
run.bold = True
add_italic_instruction(doc,
    "Write 2-3 sentences describing the workflow at a high level: what it accomplishes, who performs it today, "
    "how often it runs, and what it produces. This is the elevator pitch for the workflow — someone reading only "
    "this paragraph should understand what they are looking at."
)

add_tip(doc,
    "A good scope summary answers: What is the workflow? Who does it? How often? What does it produce? "
    "For example: 'The Quarterly Account Health Review is performed by Customer Success Managers once per "
    "quarter for each account in their portfolio. It synthesises data from CRM, support, product usage, and "
    "Slack into a scored health report with executive summary, distributed to account stakeholders.'"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 8. STEP-BY-STEP WORKFLOW MAP
# ══════════════════════════════════════════════════════════════════

doc.add_heading("8. Step-by-Step Workflow Map", level=1)

doc.add_paragraph(
    "Map every step of the workflow as you currently perform it. Write as if you are creating instructions "
    "for a new hire who has never done this work before. Be absurdly specific — the most common failure mode "
    "is abstracting away the messy details."
)

add_warning(doc,
    "\"Analyse the data\" is not a step. What data? What analysis? What are you looking for? "
    "If you can't explain it to a new hire, you can't explain it to an agent. Target a minimum of 8 steps. "
    "If you have fewer, you are almost certainly hiding complexity inside abstract step descriptions."
)

doc.add_paragraph("")

# ── Main workflow map table ───────────────────────────────────────
headers = ["#", "Step Name", "Action", "Input", "Output", "Decision Logic", "Boundary", "Error Path"]
table = add_table_with_headers(doc, headers)

# Example row from worked example
add_example_row(table, [
    "1",
    "Pull CRM account data",
    "Query CRM for account record using account ID. Retrieve account metadata including ARR, contract dates, tier, key contacts, and recent activity log.",
    "Account ID",
    "Account metadata (ARR, contract dates, tier, key contacts, recent activity log)",
    "None \u2014 data retrieval. Query CRM using account ID to retrieve a single structured record. No transformation or interpretation required.",
    "AUTOMATE",
    "If API unavailable: flag gap in report, continue with other sources. If account ID not found: halt and alert operator."
])

# Second example — a HIL step
add_example_row(table, [
    "6",
    "Assess ticket health",
    "Analyse ticket volume trends, severity distribution, resolution times vs SLA targets, and CSAT scores. Compare this quarter against prior quarter.",
    "Ticket data from Step 2, SLA benchmark targets (from config)",
    "Ticket health assessment: rating (improving / stable / degrading), SLA compliance %, CSAT trend, key issues, risk flags",
    "Criteria-based: Evaluate first response and resolution times against SLA targets by severity \u2014 P1: 30min/4hr, P2: 2hr/24hr, P3: 8hr/72hr, P4: 24hr/168hr. Flag if SLA compliance <95%. Flag if CSAT <4.0/5. Volume increase >20% QoQ = degradation.",
    "AUTOMATE (quantitative) / HIL (edge cases)",
    "If SLA config missing: use industry defaults and flag. If ticket data incomplete: calculate on available data, note gaps."
])

# Blank rows for user
for _ in range(10):
    add_blank_row(table, len(headers))

doc.add_paragraph("")

add_tip(doc,
    "For the Decision Logic column, ask the Specificity Test after each entry: 'Could someone write a prompt "
    "instruction or a config rule from this entry alone?' If no, push deeper: (1) What data do I look at? "
    "(2) What would make me choose one outcome over another?"
)

# ── Boundary tag guidance ─────────────────────────────────────────
doc.add_heading("Boundary Tag Definitions", level=2)

doc.add_paragraph(
    "Use these three tags for the Boundary column. If a step spans two boundaries, "
    "use the split format: AUTOMATE (what) / HUMAN-IN-THE-LOOP (what)."
)

bt_headers = ["Boundary Tag", "Use When", "Examples"]
bt_table = add_table_with_headers(doc, bt_headers)

add_example_row(bt_table, [
    "AUTOMATE",
    "The step is data retrieval, data transformation, or template-based generation. No judgement involved.",
    "Pull data from an API, calculate a metric, populate a report template."
])
add_example_row(bt_table, [
    "HUMAN-IN-THE-LOOP",
    "The step requires judgement, but the agent can do ~80% of the work. The human reviews and approves.",
    "Agent drafts a sentiment summary \u2014 human validates. Agent scores health \u2014 human checks edge cases."
])
add_example_row(bt_table, [
    "MANUAL",
    "The step involves relationship context, political sensitivity, or irreversible actions the agent cannot evaluate.",
    "Final review and distribution of a report to stakeholders. Decisions requiring institutional knowledge."
])

doc.add_paragraph("")

add_tip(doc,
    "The boundary should err toward more human involvement in early versions. You can push the boundary "
    "outward as trust builds. If after several cycles the human never changes the agent's output at a "
    "HIL checkpoint, that checkpoint can become fully automated. But you cannot easily add a checkpoint "
    "after users start relying on a fully automated step."
)

# ── Boundary rationale narrative ──────────────────────────────────
doc.add_heading("Boundary Rationale", level=2)

add_italic_instruction(doc,
    "For each step where the boundary assignment is not obvious (especially HIL and MANUAL steps, and any "
    "split-boundary steps), write a brief narrative explaining why you chose that boundary. This rationale "
    "is critical for Stage 4 — it determines which steps become automated nodes vs. interrupt checkpoints.\n\n"
    "Example: 'Step 5 (Slack sentiment) is HIL because Slack conversations contain sarcasm, inside jokes, and "
    "genuine complaints mixed together. The LLM can summarise and propose a sentiment reading, but the CSM "
    "knows the relationship context. The agent does the heavy lifting; the human validates.'\n\n"
    "Example: 'Step 12 (distribution) is MANUAL because once the report is sent, it is sent. The CSM needs "
    "final judgement on tone, completeness, and distribution list. This is a deliberate safety boundary.'"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 9. DATA INVENTORY
# ══════════════════════════════════════════════════════════════════

doc.add_heading("9. Data Inventory", level=1)

doc.add_paragraph(
    "Document every data source your workflow touches. This inventory becomes the tool specification "
    "for data-retrieval nodes in Stage 4 and the integration setup checklist in Stage 5."
)

di_headers = ["Data Source", "System / Platform", "Access Method", "Data Format", "Auth Required", "Rate Limits / Pagination", "Data Quality Notes", "Used by Steps"]
di_table = add_table_with_headers(doc, di_headers)

add_example_row(di_table, [
    "CRM \u2014 account metadata",
    "Salesforce",
    "REST API \u2014 query account object by account ID",
    "JSON",
    "OAuth token (connected app with refresh token flow)",
    "100K requests/day (Enterprise). Single account query = 1 call. No pagination needed.",
    "Generally reliable. Key contacts list may be stale. Contract dates are authoritative.",
    "Steps 1, 8, 9"
])

add_example_row(di_table, [
    "Slack \u2014 customer comms",
    "Slack",
    "Slack API (search.messages endpoint)",
    "JSON",
    "Bot token with search:read and channels:history scopes",
    "20 requests/min (Tier 2). Up to 100 messages/page. Expect 3-10 pages for active accounts.",
    "Inherently noisy. False positives, off-topic messages, bot messages. Signal-to-noise varies by account.",
    "Steps 4, 5"
])

for _ in range(6):
    add_blank_row(di_table, len(di_headers))

doc.add_paragraph("")

# ── Access notes narrative ────────────────────────────────────────
doc.add_heading("Access Notes and Data Quality Narrative", level=2)

add_italic_instruction(doc,
    "For each data source, describe any access challenges, data quality concerns, or known gaps that the "
    "design and build stages need to account for. This is where you capture context that does not fit in "
    "a table cell — e.g., 'The bot must be a member of all relevant Slack channels to search them. If the "
    "bot lacks access to private channels, results will be incomplete. Over-retrieve by design — filtering "
    "happens in the HIL step.' or 'CSAT scores are only present on ~35% of tickets (survey response rate). "
    "Do not treat absence of CSAT as negative.'"
)

add_warning(doc,
    "Verifying API access before starting Stage 4 is the single most common blocker in practice. "
    "Teams assume API access exists and discover during build that permissions haven't been granted, "
    "the API doesn't expose needed fields, or the API requires a platform tier upgrade."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 10. INTEGRATION REQUIREMENTS
# ══════════════════════════════════════════════════════════════════

doc.add_heading("10. Integration Requirements", level=1)

doc.add_paragraph(
    "List everything the agent will need to connect to the outside world. This table becomes the "
    "dependency checklist and environment setup guide for Stage 5 (Build)."
)

ir_headers = ["Category", "Requirement", "Version / Spec", "Status", "Owner", "Notes"]
ir_table = add_table_with_headers(doc, ir_headers)

add_example_row(ir_table, [
    "Framework",
    "LangGraph \u2014 agent orchestration",
    "LangGraph \u22650.2 (Python)",
    "Needs Setup",
    "Developer / Builder",
    "Core dependency. Workflow steps become LangGraph nodes. Install via pip install langgraph."
])

add_example_row(ir_table, [
    "LLM Provider",
    "Anthropic SDK \u2014 Claude API for synthesis, summarisation, report generation",
    "Anthropic Python SDK \u22650.39",
    "Needs Setup",
    "Developer / Builder",
    "Required for synthesis and judgement steps. API key must be provisioned and stored securely."
])

add_example_row(ir_table, [
    "Data Source API",
    "Salesforce REST API \u2014 account object read access",
    "Salesforce REST API v59.0+",
    "Needs Setup",
    "Salesforce Admin / IT",
    "Required for Step 1. Need connected app with OAuth 2.0 flow and account-read scopes."
])

add_example_row(ir_table, [
    "Authentication",
    "Credentials management \u2014 secure storage for all API keys and tokens",
    "Env vars (dev) / secrets manager (prod)",
    "Needs Setup",
    "Developer + IT",
    "Never hardcode credentials. Document required env vars in project README."
])

for _ in range(6):
    add_blank_row(ir_table, len(ir_headers))

doc.add_paragraph("")

# ── Credentials management narrative ──────────────────────────────
doc.add_heading("Credentials Management Approach", level=2)

add_italic_instruction(doc,
    "Describe how API keys, OAuth tokens, and other credentials will be managed. "
    "For example: 'For local development, credentials will be stored in a .env file loaded via python-dotenv. "
    "For production, credentials will be managed via [secrets manager]. The following environment variables "
    "are required: [list]. Salesforce OAuth tokens require automatic refresh handling.'"
)

# ── Tooling decisions narrative ───────────────────────────────────
doc.add_heading("Tooling Decisions", level=2)

add_italic_instruction(doc,
    "Document any significant tooling decisions or open questions. For example: 'Product usage data can be "
    "accessed via Mixpanel API or via data warehouse query. The Mixpanel API has strict rate limits (60 req/hr); "
    "the warehouse provides more flexible access but requires DB credentials. Decision: [pending investigation / "
    "chosen approach and why].'"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 11. CONSTRAINTS AND ASSUMPTIONS
# ══════════════════════════════════════════════════════════════════

doc.add_heading("11. Constraints and Assumptions", level=1)

doc.add_paragraph(
    "Document what must be true for this workflow to operate as designed. The assumptions you write down "
    "are rarely the ones that break you \u2014 it's the ones you don't notice. Walk your workflow map against "
    "the five diagnostic categories below to surface invisible assumptions."
)

# ── Diagnostic categories reference ──────────────────────────────
doc.add_heading("Diagnostic Categories", level=2)

doc.add_paragraph(
    "For each category, ask the diagnostic question and check whether your workflow silently depends on "
    "the answer being a specific value."
)

dc_headers = ["Category", "Diagnostic Question"]
dc_table = add_table_with_headers(doc, dc_headers)

dc_rows = [
    ("Entity Cardinality", "Does any step operate on a single entity that could, in production, be multiple? (One product per account, one contact per deal, one config per environment.)"),
    ("Temporal / Frequency", "Does any step assume data is current as of a specific window, or that the workflow runs at a specific cadence?"),
    ("Access / Permissions", "Does any step assume the agent can read or write data that might require elevated permissions, user consent, or manual export?"),
    ("Stability / Change Rate", "Does any step depend on a schema, scoring framework, or business rule that could change mid-quarter?"),
    ("Processing Model", "Does any step assume sequential processing that could encounter parallel, batched, or out-of-order inputs?"),
]
for cat, question in dc_rows:
    row = dc_table.add_row()
    row.cells[0].text = cat
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
    row.cells[1].text = question
    for p in row.cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)

doc.add_paragraph("")

# ── Main constraints table ────────────────────────────────────────
doc.add_heading("Constraints and Assumptions Register", level=2)

ca_headers = ["Category", "Assumption Statement", "What Breaks If Wrong", "Verified?", "Impact If Wrong", "Mitigation / Fallback"]
ca_table = add_table_with_headers(doc, ca_headers)

add_example_row(ca_table, [
    "Entity Cardinality",
    "Each account has exactly one active product. The health scoring workflow evaluates a single product's usage metrics per account.",
    "If multi-product: health score conflates products \u2014 healthy Product A masks failing Product B. Steps 7 and 9 need product-level sub-steps.",
    "Needs investigation",
    "Stage 3 (workflow map), Stage 4 (state schema needs product loop), Stage 5 (scoring rubric needs per-product dimensions)",
    "If multi-product accounts exist: (1) add product enumeration step, (2) loop scoring per product, (3) add roll-up step. Or scope v1 to single-product accounts only."
])

add_example_row(ca_table, [
    "Access / Permissions",
    "All data sources have programmatic API access with read permissions granted to the agent's service account.",
    "If any source requires manual export, AUTOMATE boundary tags for those steps change to MANUAL, altering the workflow's automation profile.",
    "Needs investigation",
    "Stage 3 (boundary revision), Stage 4 (manual steps become wait states), Stage 5 (build scope shrinks)",
    "Verify API access for each source before Stage 4. If unavailable: (1) check if API can be enabled, (2) check for scheduled export workaround, (3) reclassify step as MANUAL."
])

add_example_row(ca_table, [
    "Temporal / Frequency",
    "The workflow runs quarterly. Metrics are compared QoQ using aligned quarterly boundaries. Data sources refresh within 72 hours of quarter-end.",
    "If run mid-quarter: partial vs full quarter comparisons produce misleading trends. If data sources have different lag: metrics from same date range represent different time windows.",
    "No",
    "Stage 3 (date range params need freshness checks), Stage 5 (retrieval nodes need configurable dates)",
    "Define explicit date range params at workflow start. Add data freshness check per source. If any source lags >72hrs, flag in report header."
])

for _ in range(6):
    add_blank_row(ca_table, len(ca_headers))

doc.add_paragraph("")

add_warning(doc,
    "The entity cardinality assumption (single product per account, single owner per entity) is the "
    "assumption most likely to be missed. It is invisible because you implicitly assume it while mapping "
    "the workflow. Surface it explicitly and verify against your actual data before proceeding to Stage 4."
)

# ── Diagnostic reasoning narrative ────────────────────────────────
doc.add_heading("Diagnostic Reasoning", level=2)

add_italic_instruction(doc,
    "For each assumption in the register above, write a brief narrative explaining how you surfaced it "
    "and what breaks if it turns out to be wrong. This narrative is more valuable than the table entry "
    "because it captures the reasoning that led to the assumption — not just the assumption itself.\n\n"
    "Be specific about the assumed value, not just the category. Bad: 'The agent operates per-account.' "
    "Good: 'The agent operates per-account, processing one account at a time. Each account has exactly one "
    "active product — if an account has multiple products, the health scoring step would need to evaluate "
    "and weight each product separately, which this workflow does not handle.'\n\n"
    "For each entry, state: (1) which diagnostic category surfaced it, (2) the specific value you are "
    "assuming, (3) what downstream stages are affected if the assumption is wrong, and (4) whether you "
    "have verified it or it needs investigation."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 12. COMPLETION CHECKLIST
# ══════════════════════════════════════════════════════════════════

doc.add_heading("12. Completion Checklist", level=1)

doc.add_paragraph(
    "Review your completed document against this checklist. Every item should be ticked before "
    "handing this document to the person responsible for Stage 4 (Design)."
)

checklist_items = [
    "Workflow Overview is complete: workflow name, Stage 2 reference, and scope summary are filled in.",
    "Workflow Map has a minimum of 8 steps (if fewer, break abstract steps down further).",
    "Every step has an Action description specific enough for a new hire to follow.",
    "Every step has Input and Output columns filled in.",
    "Decision Logic column is completed for every step, using the correct type (None, Criteria-based, or Contextual judgement).",
    "Decision Logic entries pass the Specificity Test: someone could write a prompt instruction or config rule from each entry alone.",
    "Every step has a Boundary tag (AUTOMATE, HUMAN-IN-THE-LOOP, MANUAL, or a documented split).",
    "Boundary Rationale narrative is written for all HIL, MANUAL, and split-boundary steps.",
    "Error Path is documented for every step (what happens when the step fails).",
    "Critical dependencies (steps where failure halts the entire workflow) are flagged in Constraints and Assumptions.",
    "Data Inventory is complete for every data source the workflow touches.",
    "Access Notes narrative documents known data quality issues and access challenges.",
    "Integration Requirements table lists all frameworks, APIs, credentials, and output tooling.",
    "Credentials Management approach is documented.",
    "Constraints and Assumptions register has been walked against all five diagnostic categories.",
    "Each assumption states a specific value, not just a category name.",
    "Assumptions are marked as Verified or Needs Investigation — none are left unmarked.",
    "Diagnostic Reasoning narrative explains how each assumption was surfaced and what breaks if wrong.",
    "All italic instructional text has been replaced with your own content.",
]

for item in checklist_items:
    add_checkbox_item(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 13. NOTES & OPEN QUESTIONS
# ══════════════════════════════════════════════════════════════════

doc.add_heading("13. Notes & Open Questions", level=1)

add_italic_instruction(doc,
    "Use this section to capture anything you are unsure about, questions that need answers before "
    "Stage 4 can begin, or observations that do not fit neatly into the sections above. Open questions "
    "are expected — it is better to surface uncertainty here than to hide it in an assumption.\n\n"
    "Examples:\n"
    "- 'Need to confirm whether the Mixpanel API exposes feature adoption data or if we need a warehouse query.'\n"
    "- 'The health scoring rubric was last updated Q3 2024 — need to verify it is still current.'\n"
    "- 'Unsure whether the Slack bot has been added to the #enterprise-accounts private channel.'"
)

doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

# ══════════════════════════════════════════════════════════════════
# 14. NEXT STEPS
# ══════════════════════════════════════════════════════════════════

doc.add_heading("14. Next Steps", level=1)

doc.add_paragraph(
    "Once this document is complete and all checklist items are verified:"
)

next_steps = [
    ("Resolve open questions.", "Any item marked 'Needs Investigation' in the Constraints and Assumptions register should be verified before proceeding. Unverified assumptions become design risks in Stage 4."),
    ("Hand off to Stage 4: Design.", "The completed Scope Document is the primary input for Stage 4, where the workflow map is translated into a LangGraph agent architecture. The designer will use your step-by-step map to define nodes and edges, your boundary tags to place interrupt checkpoints, your data inventory to specify tools, and your constraints to shape the state schema."),
    ("Retain this document as a living reference.", "As the agent is built and tested in Stages 5 and 6, you may discover that some assumptions were wrong or that the workflow needs adjustment. Update this document to reflect what you learn — it remains the authoritative record of what the agent is designed to do and why."),
]

for i, (title, desc) in enumerate(next_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {title} ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph("")

add_tip(doc,
    "The quality of Stage 4 (Design) is entirely determined by the quality of this Scope Document. "
    "Time invested in capturing decision logic, boundary rationale, and constraints here saves "
    "significantly more time during design and build."
)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), "Stage 3 - Scope - Template.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
