#!/usr/bin/env python3
"""Build Stage 4 - Design - Template.docx

Generates a professional Word document template for Stage 4 of the
agentic workflow framework, matching the styling conventions established
by the Stage 3 template.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour constants (matching Stage 3) ─────────────────────────────
NAVY = RGBColor(0x1B, 0x3A, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x66, 0x66, 0x66)
HEADER_FILL = "1B3A5C"
EXAMPLE_FILL = "D6E4F0"
TIP_BG = "E8F4E8"
WARNING_BG = "FFF3E0"


# ── Helper functions ────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_header_row(table, headers, font_size=Pt(9)):
    """Style the first row of a table as a header row."""
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = font_size
        set_cell_shading(cell, HEADER_FILL)


def add_example_row(table, row_idx, values, font_size=Pt(8)):
    """Add an example row with light blue shading."""
    for i, text in enumerate(values):
        cell = table.rows[row_idx].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = font_size
        run.italic = True
        set_cell_shading(cell, EXAMPLE_FILL)


def add_empty_rows(table, count, col_count, font_size=Pt(8)):
    """Add empty rows for user to fill in."""
    for _ in range(count):
        row = table.add_row()
        for i in range(col_count):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run("")
            run.font.size = font_size


def add_tip(doc, text):
    """Add an indented TIP callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Emu(360045)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run("TIP: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)


def add_warning(doc, text):
    """Add an indented WARNING callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Emu(360045)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run("WARNING: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)


def add_italic_instruction(doc, text):
    """Add an italic instructional paragraph the user replaces."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    return p


def add_body(doc, text, size=None):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    if size:
        for run in p.runs:
            run.font.size = size
    return p


def add_bullet(doc, text):
    """Add a bullet list item."""
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def set_table_width(table, width_inches=6.5):
    """Set table to a fixed width."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


# ── Build the document ──────────────────────────────────────────────

doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════

# Spacer lines for vertical centering
for _ in range(6):
    doc.add_paragraph("")

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Stage 4: Design")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = NAVY

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Agent Design Document")
run.font.size = Pt(16)
run.font.color.rgb = NAVY

doc.add_paragraph("")

# Workflow Name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Workflow Name: ")
run.font.size = Pt(12)
run = p.add_run("________________________________________")
run.font.size = Pt(12)
run.font.color.rgb = GREY

doc.add_paragraph("")

# Date / Author / Version
for label in ["Date:  ", "Author:  ", "Version:  "]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(11)
    run = p.add_run("________________________")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

# Page break
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. DOCUMENT PURPOSE
# ════════════════════════════════════════════════════════════════════

doc.add_heading("1. Document Purpose", level=1)

add_body(doc,
    "This document translates your Workflow Scope Document (Stage 3) into an "
    "engineering blueprint for an AI agent. Where the scope document describes "
    "what the workflow does in operational language, this design document "
    "specifies how an agent will execute it \u2014 which units of work (actions) the "
    "agent performs, what data flows between them, where humans review and "
    "intervene, and how the agent handles failures. Think of it as the floor "
    "plan before construction: every architectural decision is made here so "
    "that Stage 5 (Build) can focus on execution."
)

add_body(doc,
    "The design document serves three audiences. For the builder (developer or "
    "technical implementer), it provides enough detail to implement each action "
    "independently given only this document and the memory schema. For the "
    "workflow owner (CSM, analyst, or subject-matter expert), it confirms that "
    "the scope decisions survive translation into agent architecture \u2014 every "
    "automation boundary, every human checkpoint, every error strategy traces "
    "back to a scope step. For reviewers and future maintainers, it documents "
    "the consolidation rationale: why 12 scope steps became 6 actions, why 7 "
    "HIL candidates became 2 HIL checkpoints, and why each design choice was "
    "made."
)

add_body(doc,
    "Once completed, this document feeds directly into Stage 5 (Build). The "
    "flow structure becomes the platform\u2019s flow construct. The memory schema "
    "becomes the typed memory/state object used by the chosen framework. The "
    "action specifications become the individual action implementations. The "
    "HIL interaction designs become the pause-and-resume implementations. "
    "Nothing in this document requires code \u2014 but everything in it will be "
    "implemented as code in the next stage."
)

# ════════════════════════════════════════════════════════════════════
# 2. SCOPE
# ════════════════════════════════════════════════════════════════════

doc.add_heading("2. Scope", level=1)

doc.add_heading("What This Document Covers", level=2)

add_body(doc,
    "This document covers the complete agent architecture for a single "
    "workflow, including:"
)

for item in [
    "Mapping of every scope step to a flow action, with consolidation rationale for grouping and conversion decisions",
    "Flow structure \u2014 actions, edges (unconditional and conditional), and control flow from start to end",
    "Memory schema \u2014 every typed field the flow needs, with derivation tracing each field to a scope step",
    "Action specifications \u2014 purpose, tools, logic, prompt patterns, and error handling for each action",
    "Human-in-the-loop interaction design \u2014 HIL pattern (feedback-as-context or feedback-as-control-flow), surfaced data, and resume format for each checkpoint",
    "Error handling design \u2014 failure modes, detection methods, and response strategies for every action",
]:
    add_bullet(doc, item)

doc.add_heading("What This Document Does NOT Cover", level=2)

for item in [
    "Workflow discovery or selection rationale (Stage 2: Select \u2014 assumed complete)",
    "Operational workflow detail \u2014 step-by-step actions, data sources, decision logic (Stage 3: Scope \u2014 assumed complete)",
    "Implementation code, prompt text, or test cases (Stage 5: Build)",
    "Deployment, monitoring, or production operations (Stage 6: Evaluate & Iterate)",
    "Multi-workflow orchestration or batch processing",
]:
    add_bullet(doc, item)

# ════════════════════════════════════════════════════════════════════
# 3. OBJECTIVES
# ════════════════════════════════════════════════════════════════════

doc.add_heading("3. Objectives", level=1)

add_body(doc, "By completing this document, you will have accomplished the following:")

objectives = [
    "Mapped every AUTOMATE and HUMAN-IN-THE-LOOP scope step to a flow action, documenting why steps were grouped, converted, or kept separate.",
    "Consolidated HIL checkpoints from the initial candidate list down to the minimum needed for effective human oversight, with written rationale for each consolidation decision.",
    "Drawn a flow structure that has a clear path from start to end for every possible execution, including conditional paths for HIL decision points.",
    "Defined a typed memory schema where every field traces to a specific scope step output, data inventory entry, or HIL interaction pattern.",
    "Written action specifications detailed enough that a developer can implement each action independently given only the spec and the memory schema.",
    "Designed HIL interactions that specify what the human sees, what they provide, and how their input affects flow execution (routing vs. context enrichment).",
    "Documented error handling strategies for every identified failure mode, choosing deliberately between error-marker-and-continue, retry, escalate, or abort.",
]

for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    run2 = p.add_run(obj)

# ════════════════════════════════════════════════════════════════════
# 4. INPUTS REQUIRED
# ════════════════════════════════════════════════════════════════════

doc.add_heading("4. Inputs Required", level=1)

add_body(doc,
    "Before starting this document, ensure you have the following:"
)

inputs = [
    ("Workflow Scope Document (Stage 3)", "The completed scope document for this workflow, including the workflow map (step-by-step table), data inventory, boundary assignments, and assumptions register. This is your primary source \u2014 every design decision traces back to it."),
    ("Access to the Stage 3 author or SME", "You will need to ask clarifying questions about decision logic, edge cases, and boundary assignments. If you are the Stage 3 author, you can proceed independently."),
    ("Familiarity with agent-framework core concepts", "You do not need to write code, but you need to understand actions, edges, memory, and the pause-and-resume pattern. See the Quick Reference table in Stage 4 methodology."),
    ("Scoring rubrics or decision frameworks (if applicable)", "If the workflow uses structured scoring criteria (e.g., health scoring rubrics, SLA thresholds, rating frameworks), have these available \u2014 they inform the prompt pattern and config design in action specifications."),
]

for title, desc in inputs:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}. ")
    run.bold = True
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════
# 5. OUTPUTS PRODUCED
# ════════════════════════════════════════════════════════════════════

doc.add_heading("5. Outputs Produced", level=1)

add_body(doc,
    "This document produces the following artifacts that Stage 5 (Build) consumes directly:"
)

outputs = [
    ("Scope-to-Action Mapping Table", "Documents which scope steps map to which flow actions, with consolidation rationale. Used in Stage 5 to verify implementation completeness and in Stage 6 to trace evaluation findings back to design decisions."),
    ("Flow Structure", "Visual representation of actions, edges, and control flow. Becomes the platform\u2019s flow construct in Stage 5."),
    ("Memory Schema", "Typed memory shape with every field the flow needs. Becomes the memory / state definition in Stage 5."),
    ("Action Specifications", "Detailed spec cards (purpose, tools, logic, prompt pattern, error handling) for each action. Each spec becomes a concrete action implementation in Stage 5."),
    ("HIL Interaction Design", "Pattern choice, surfaced data, and resume format for each checkpoint. Becomes the pause-and-resume implementation in Stage 5."),
    ("Error Handling Design", "Failure modes, detection methods, and response strategies. Implemented as try/except blocks, retry loops, and escalation logic in Stage 5."),
]

for title, desc in outputs:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}. ")
    run.bold = True
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════
# 6. HOW TO USE THIS DOCUMENT
# ════════════════════════════════════════════════════════════════════

doc.add_heading("6. How to Use This Document", level=1)

add_body(doc,
    "Work through this document sequentially \u2014 each section builds on the previous one. "
    "The sections are ordered to match the natural design workflow: map scope steps to "
    "actions first, then draw the flow structure, then define the memory, then write action "
    "specs, then design HIL interactions, then plan error handling."
)

steps = [
    "Read the guidance text in each section to understand what is expected and why it matters.",
    "Fill in the tables using the column structure provided. Each table includes detailed example rows (in blue) showing the expected depth and format \u2014 match that level of detail.",
    "Write the narrative sections where indicated by italic instructional text. Replace the italic text with your own content.",
    "Review your work against the Completion Checklist at the end of the document.",
    "Capture any unresolved questions in the Notes & Open Questions section.",
]

for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(step)

add_tip(doc,
    "If you are not a developer, this is where you collaborate with one. Your scope "
    "document is in operational language \u2014 this design document translates it into "
    "engineering language. You make the decisions about structure, data flow, and human "
    "involvement; the developer validates technical feasibility."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 7. SCOPE-TO-ACTION MAPPING
# ════════════════════════════════════════════════════════════════════

doc.add_heading("7. Scope-to-Action Mapping", level=1)

doc.add_heading("7.1 Consolidation Rationale", level=2)

add_body(doc,
    "Before filling in the mapping table, walk through your Workflow Scope Document "
    "and apply the mapping rules: (1) each AUTOMATE step becomes an action, (2) each "
    "HUMAN-IN-THE-LOOP step is a candidate for an HIL checkpoint action, (3) MANUAL steps "
    "are outside the flow, (4) group tightly coupled steps into a single action if "
    "they share the same tool and logic pattern, (5) identify steps that can run in "
    "parallel, and (6) define conditional edges where the next step depends on the "
    "outcome."
)

add_body(doc,
    "After the initial mapping pass, apply the two consolidation principles to reduce "
    "HIL checkpoints: Principle 1 (merge related reviews into a single checkpoint when "
    "the reviewer would naturally evaluate the outputs together) and Principle 2 (let "
    "upstream checkpoints convert downstream HIL steps to automated when the human has "
    "already validated the inputs)."
)

add_italic_instruction(doc,
    "Write 2\u20133 paragraphs explaining your consolidation decisions. How many scope steps "
    "did you start with? How many flow actions did you consolidate them into? How many "
    "HIL candidates did you reduce to how many HIL checkpoints? For each consolidation, "
    "explain the reasoning \u2014 which rule or principle applied and why. Reference specific "
    "scope step numbers."
)

add_warning(doc,
    "If your initial mapping produces more than 3\u20134 HIL checkpoints, revisit your "
    "consolidation. Excessive checkpoints degrade the agent's usability \u2014 the human spends "
    "more time reviewing than the agent saves. Apply the consolidation principles "
    "systematically before finalising."
)

doc.add_heading("7.2 Scope-to-Action Mapping Table", level=2)

add_body(doc,
    "Record every scope step and its mapping decision. One row per scope step. "
    "Add additional rows for any new actions that emerge from consolidation "
    "(e.g., consolidated HIL checkpoint actions that don't map 1:1 to a scope step). "
    "The four example rows below demonstrate the most common mapping patterns you "
    "will encounter: simple AUTOMATE grouping, split-boundary decomposition, "
    "consolidated HIL checkpoint creation, and MANUAL step exclusion."
)

# Mapping table — columns match the CSV
mapping_headers = [
    "Scope Step\nNumber",
    "Scope Step\nName",
    "Boundary Tag\n(from Stage 3)",
    "Flow\nAction",
    "Consolidation\nRule Applied",
    "Consolidation Rationale",
    "Notes",
]

table = doc.add_table(rows=1, cols=7)
table.style = "Table Grid"
add_header_row(table, mapping_headers, Pt(8))

# Example row 1 — Simple AUTOMATE grouping (original)
ex_row = table.add_row()
add_example_row(table, 1, [
    "1",
    "Pull CRM data",
    "AUTOMATE",
    "gather_data",
    "Rule 1 (AUTOMATE \u2192 action) + Rule 4 (group tightly coupled)",
    "All data retrieval steps (1\u20134) share the same logic pattern: API call with account_id + "
    "date range. Different data sources but identical action behaviour. Grouping keeps the "
    "flow simple while internal parallelism handles concurrency.",
    "API calls within this action are independent and can be parallelised internally",
], Pt(7))

# Example row 2 — Split-boundary step (AUTOMATE/HIL decomposition)
ex_row2 = table.add_row()
add_example_row(table, 2, [
    "6",
    "Assess ticket health",
    "AUTOMATE / HIL (split boundary)",
    "analyse_health",
    "Rule 1 (AUTOMATE component) + Principle 1 (HIL component merged into single checkpoint)",
    "Split-boundary step: the quantitative analysis (calculate volume trends / severity "
    "distribution / MTTR vs SLA / CSAT) is automated as processing logic inside "
    "analyse_health. The HIL component (interpreting edge cases and validating scores) "
    "is consolidated with Steps 5 and 7 into the hil_review_analysis checkpoint. "
    "Splitting follows the methodology's instruction to decompose split-boundary steps "
    "before applying rules.",
    "SLA thresholds and scoring criteria come from an external scoring rubric config \u2014 "
    "not hardcoded in the action",
], Pt(7))

# Example row 3 — Consolidated HIL checkpoint (new node, not 1:1 with scope step)
ex_row3 = table.add_row()
add_example_row(table, 3, [
    "\u2014",
    "\u2014",
    "\u2014",
    "hil_review_analysis",
    "Principle 1 (merge related reviews into single checkpoint)",
    "This action is a new consolidated HIL checkpoint for Steps 5\u20137. Rather than three "
    "separate checkpoints (one per analysis dimension), the flow pauses once and surfaces "
    "all three analyses together. The reviewer gets full context to cross-reference "
    "findings and can correct any dimension in a single interaction. This eliminates "
    "context-switching overhead and produces higher-quality feedback.",
    "Uses Feedback-as-Context pattern: human provides freeform text stored in "
    "analysis_human_feedback. Unconditional edge to synthesise.",
], Pt(7))

# Example row 4 — MANUAL step mapped to 'Outside graph'
ex_row4 = table.add_row()
add_example_row(table, 4, [
    "12",
    "Final review & distribution",
    "MANUAL",
    "Outside flow",
    "Rule 3 (MANUAL steps are outside the flow)",
    "MANUAL steps happen after the flow completes. The flow ends when "
    "hil_review_report approves the output. Final edits, distribution list decisions, "
    "and the actual send all happen outside the flow invocation. The flow boundary ends at "
    "the last point where the agent adds value.",
    "Step 12 was tagged MANUAL in Stage 3 because distribution involves relationship "
    "decisions and channel choices that require human judgement",
], Pt(7))

# Add empty rows
add_empty_rows(table, 12, 7, Pt(8))

add_tip(doc,
    "Split-boundary steps (tagged both AUTOMATE and HIL) should be decomposed before "
    "applying the mapping rules. The AUTOMATE component becomes processing logic inside "
    "an action; the HIL component becomes a candidate for consolidation with other checkpoints. "
    "See example row 2 above for this pattern in practice."
)

add_tip(doc,
    "Consolidated HIL checkpoint actions (like hil_review_analysis in example row 3) do not "
    "map 1:1 to a scope step \u2014 they are new actions created by applying Principle 1. Add these "
    "as additional rows with \u2018\u2014\u2019 in the Scope Step Number and Name columns to show they "
    "emerged from consolidation."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 8. FLOW STRUCTURE
# ════════════════════════════════════════════════════════════════════

doc.add_heading("8. Flow Structure", level=1)

doc.add_heading("8.1 Pattern Selection", level=2)

add_body(doc,
    "Identify which design pattern(s) your flow follows. Most workflows combine patterns. "
    "The four core patterns are: (1) Sequential Pipeline with HIL Checkpoints, "
    "(2) Parallel Fan-Out / Fan-In, (3) Router, and (4) Iterative Refinement Loop."
)

add_italic_instruction(doc,
    "State which pattern(s) your flow uses and why. Reference the structure of your "
    "scope document \u2014 e.g., 'The workflow is fundamentally linear with two review gates, "
    "so the dominant pattern is Sequential Pipeline with HIL Checkpoints. Data gathering "
    "uses internal parallelism within a single action rather than flow-level fan-out because "
    "the parallel branches share the same logic pattern.' If you use sub-flow composition, "
    "explain why the flow warrants decomposition (8+ actions, distinct phases, reusability)."
)

doc.add_heading("8.2 Flow Diagram", level=2)

add_body(doc,
    "Draw your flow structure below. This can be a Mermaid diagram, an ASCII diagram, or a "
    "hand-drawn sketch (photograph and paste). The diagram must show: all actions (including "
    "HIL checkpoint actions), all edges (unconditional and conditional), start and end, "
    "and labels on conditional edges indicating the routing condition."
)

add_body(doc,
    "Below is an example from the worked example (Quarterly Account Health Review) showing "
    "a Sequential Pipeline with HIL Checkpoints pattern, including both unconditional edges "
    "and conditional routing from a Feedback-as-Control-Flow checkpoint:"
)

add_italic_instruction(doc,
    "Replace this example with your own flow structure diagram.\n\n"
    "flow TD\n"
    "    START([START]) --> gather_data[gather_data]\n"
    "    gather_data --> analyse_health[analyse_health]\n"
    "    analyse_health --> hil_review_analysis{{HIL: review analysis}}\n"
    "    hil_review_analysis -->|human confirms/edits| synthesise[synthesise]\n"
    "    synthesise --> generate_report[generate_report]\n"
    "    generate_report --> hil_review_report{{HIL: review report}}\n"
    "    hil_review_report -->|approved| END([END])\n"
    "    hil_review_report -->|revise| generate_report\n"
    "    hil_review_report -->|rework| analyse_health"
)

# Spacer for diagram
for _ in range(3):
    doc.add_paragraph("")

add_warning(doc,
    "If you cannot draw the flow clearly, your scope probably needs more decomposition. "
    "Spend 15 minutes with the diagram before proceeding \u2014 discovering structural problems "
    "here is dramatically cheaper than discovering them mid-build."
)

doc.add_heading("8.3 Flow Notes", level=2)

add_italic_instruction(doc,
    "Add any explanatory notes about your flow decisions. Why is this a flat flow "
    "vs. sub-flows? Why are certain steps consolidated into one action rather than separate "
    "actions? Why are there two checkpoints rather than one (or three)? These notes help "
    "reviewers and future maintainers understand the design intent."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 9. MEMORY SCHEMA
# ════════════════════════════════════════════════════════════════════

doc.add_heading("9. Memory Schema", level=1)

doc.add_heading("9.1 Field Derivation Notes", level=2)

add_body(doc,
    "Memory fields come from three sources in the scope document. Walk through each "
    "systematically: (1) Input fields \u2014 what the flow needs to start (from the Input "
    "column of your first scope steps), (2) Gathered and intermediate fields \u2014 what each "
    "action produces (from the Output column of scope steps), and (3) HIL feedback fields \u2014 "
    "what the human provides at each checkpoint."
)

add_italic_instruction(doc,
    "Explain how you derived your memory fields. For each category (inputs, gathered data, "
    "analysis outputs, HIL feedback, synthesis, final output), describe which scope step "
    "outputs and data inventory entries each field traces to. If a field uses a reducer "
    "(e.g., an append reducer), explain why \u2014 which actions write to it "
    "and why accumulation is needed."
)

add_tip(doc,
    "If you cannot point to a specific scope document row or data inventory entry as the "
    "source for a field, either the scope document is missing something or the field does "
    "not belong in the schema."
)

doc.add_heading("9.2 Memory Schema Table", level=2)

add_body(doc,
    "Document every field in your memory schema. One row per field. The example rows "
    "below demonstrate all four field categories: Input, Gathered Data (with reducer), "
    "Analysis (intermediate output), and Human Feedback."
)

# Memory schema table — columns match the CSV
schema_headers = [
    "Field Name",
    "Type",
    "Category",
    "Source\n(which action writes)",
    "Consumed By\n(which actions read)",
    "Reducer",
    "Derived From\n(scope step Output)",
    "Notes",
]

table = doc.add_table(rows=1, cols=8)
table.style = "Table Grid"
add_header_row(table, schema_headers, Pt(7))

# Example row 1 — Input field
ex_row = table.add_row()
add_example_row(table, 1, [
    "account_id",
    "string",
    "Input",
    "Caller (flow invocation)",
    "gather_data; analyse_health; synthesise; generate_report",
    "",
    "Scope Step 1 Input: account identifier used to query all data sources",
    "Required input \u2014 the caller must supply this when invoking the flow",
], Pt(7))

# Example row 2 — Gathered Data with reducer
ex_row2 = table.add_row()
add_example_row(table, 2, [
    "tickets",
    "list[dict] (append reducer)",
    "Gathered Data",
    "gather_data",
    "analyse_health",
    "append (list accumulation)",
    "Scope Step 2 Output: support ticket records from the support API",
    "Reducer supports future fan-out if data gathering is split into parallel sub-actions",
], Pt(7))

# Example row 3 — Analysis (intermediate non-reducer field)
ex_row3 = table.add_row()
add_example_row(table, 3, [
    "ticket_analysis",
    "dict",
    "Analysis",
    "analyse_health",
    "synthesise; generate_report; hil_review_analysis (surfaced to human)",
    "",
    "Scope Step 6 Output: ticket health score with trend and findings",
    "JSON structure: {score, trend, volume_vs_previous_quarter, key_issues, sla_compliance, details}. "
    "Score set to 'Unknown' if upstream tickets field contains error marker.",
], Pt(7))

# Example row 4 — HIL feedback field
ex_row4 = table.add_row()
add_example_row(table, 4, [
    "analysis_human_feedback",
    "string (optional)",
    "Human Feedback",
    "hil_review_analysis",
    "synthesise; analyse_health (on rework)",
    "",
    "HIL checkpoint 1: freeform human corrections / additions about the analyses",
    "Feedback-as-Context pattern. None on first pass.",
], Pt(7))

add_empty_rows(table, 15, 8, Pt(7))

doc.add_heading("9.3 Memory Schema Pseudocode", level=2)

add_body(doc,
    "Translate the table above into a platform-neutral memory schema sketch. This "
    "sketch becomes the typed memory/state object in whichever framework your team "
    "has chosen for Stage 5. Include comments grouping fields by category (Input, "
    "Gathered Data, Analysis, Human Feedback, Synthesis, Output)."
)

add_italic_instruction(doc,
    "Write your memory schema here as pseudocode. Follow this shape:\n\n"
    "memory YourWorkflowMemory:\n"
    "    // --- Input ---\n"
    "    field_name: string\n\n"
    "    // --- Gathered Data ---\n"
    "    data_field: dict\n"
    "    list_field: list[dict]  // append reducer for accumulation\n\n"
    "    // --- Analysis ---\n"
    "    analysis_result: dict\n\n"
    "    // --- Human Feedback ---\n"
    "    feedback_field: string (optional)  // feedback-as-context\n"
    "    decision_field: string (optional)  // feedback-as-control-flow\n\n"
    "    // --- Synthesis ---\n"
    "    risks: list[dict]\n"
    "    health_score: string\n\n"
    "    // --- Output ---\n"
    "    final_output: string\n\n"
    "Your chosen agent framework (see Stage 4: Choose a Platform) will provide the "
    "concrete construct (typed dict, dataclass, schema object, etc.) for representing "
    "this memory."
)

add_warning(doc,
    "Designing memory as an afterthought is a common mistake. The memory schema is the "
    "contract between your actions \u2014 define it before writing action specs, and ensure every "
    "action spec references memory fields that actually exist in this schema."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 10. ACTION SPECIFICATIONS
# ════════════════════════════════════════════════════════════════════

doc.add_heading("10. Action Specifications", level=1)

add_body(doc,
    "Write one specification card per action in your flow structure. Each card must "
    "contain enough detail that a developer can implement the action independently given "
    "only the card and the memory schema. Derive each field from the scope document: "
    "Purpose from the Output column, Tools from the Data Inventory, Logic from the "
    "Decision Logic column, and Prompt Pattern from the criteria and thresholds in "
    "Decision Logic."
)

add_tip(doc,
    "If you cannot state an action's purpose in one sentence, the action is doing too much \u2014 "
    "split it. An action that 'gathers data, analyses it, and drafts a summary' is three actions."
)

doc.add_heading("10.1 Action Specification Table", level=2)

add_body(doc,
    "The table below captures all action specifications in a structured format. "
    "One row per action. For complex actions, the Logic and Prompt Pattern cells "
    "may contain multiple numbered steps. The three example rows demonstrate "
    "a Data Gathering action, an LLM Analysis action, and a HIL Checkpoint action."
)

# Action specifications table — columns match the CSV
node_headers = [
    "Action\nName",
    "Purpose",
    "Action\nType",
    "Tools",
    "Logic",
    "Prompt Pattern",
    "Error\nHandling",
    "Input Memory\nFields",
    "Output Memory\nFields",
    "Edge\nType",
    "Notes",
]

table = doc.add_table(rows=1, cols=11)
table.style = "Table Grid"
add_header_row(table, node_headers, Pt(7))

# Example row 1 — data gathering action
ex_row = table.add_row()
add_example_row(table, 1, [
    "gather_data",
    "Retrieve all raw data for the account from external APIs",
    "Data Gathering",
    "CRM API wrapper; Support API wrapper; Usage data API; Slack search",
    "1. Read account_id and quarter from memory. "
    "2. Call each API with account_id and date range. "
    "3. Store each response in corresponding memory field. "
    "All API calls are independent \u2014 parallelise internally.",
    "N/A \u2014 no LLM calls. Deterministic data retrieval.",
    "If any API fails: store error marker dict in memory field, continue with remaining sources. "
    "If source returns empty: store empty result (not error marker).",
    "account_id; quarter",
    "crm_data; tickets; usage_data; slack_threads",
    "Unconditional \u2192 analyse_health",
    "Four API calls parallelised internally. Use fan-out at flow level only when branches have different logic.",
], Pt(6))

# Example row 2 — LLM analysis action
ex_row2 = table.add_row()
add_example_row(table, 2, [
    "analyse_health",
    "Score ticket health and product engagement from raw data",
    "LLM Analysis",
    "LLM \u2014 three separate calls, one per dimension",
    "1. Check upstream data for error markers; skip dimensions with errors. "
    "2. Ticket: calculate volume trends, MTTR vs SLA, pass to LLM for scoring. "
    "3. Usage: calculate DAU/MAU trends, adoption rates, pass to LLM. "
    "4. Sentiment: pass Slack threads to LLM for theme identification. "
    "5. On rework: include rework_instructions in all prompts.",
    "Three prompts, each receiving account_name + quarter. "
    "Ticket: raw data + SLA thresholds from config. JSON output: {score, trend, key_issues, ...}. "
    "Usage: metrics + adoption benchmarks from config. JSON output: {score, trend, ...}. "
    "Sentiment: Slack threads, qualitative. JSON output: {score, summary, key_threads}.",
    "Error marker upstream \u2192 skip dimension, score 'Unknown'. "
    "Malformed JSON \u2192 retry with stricter prompt (up to 2 retries). "
    "Retries exhausted \u2192 escalate via an HIL checkpoint. "
    "Dimensions fail independently.",
    "crm_data; tickets; usage_data; slack_threads; rework_instructions (on rework)",
    "ticket_analysis; usage_analysis; sentiment_analysis",
    "Unconditional \u2192 hil_review_analysis",
    "Consolidates scope steps 5\u20137. Three LLM calls can be parallelised internally.",
], Pt(6))

# Example row 3 — HIL Checkpoint action
ex_row3 = table.add_row()
add_example_row(table, 3, [
    "hil_review_report",
    "Present complete report to human for final review with approve/revise/rework options",
    "HIL Checkpoint",
    "None \u2014 the HIL checkpoint pauses execution",
    "1. Format report package: report_markdown + exec_summary + health_score + risks + opportunities. "
    "2. Pause at the HIL checkpoint with the formatted payload and decision options. "
    "3. On resume: store decision in report_review_decision. "
    "4. Route instructions to correct memory field based on decision.",
    "N/A \u2014 no LLM calls. Surfaces memory data to human and processes structured response.",
    "Memory persists via the chosen framework's checkpointing layer \u2014 no timeout crash. "
    "Routing function validates decision value; raises an error on unrecognised value.",
    "report_markdown; exec_summary; health_score; risks; opportunities",
    "report_review_decision; report_human_feedback (revise); rework_instructions (rework)",
    "Conditional: approved \u2192 END; revise \u2192 generate_report; rework \u2192 analyse_health",
    "Feedback-as-Control-Flow pattern. Three decision paths give reviewer proportional control.",
], Pt(6))

add_empty_rows(table, 8, 11, Pt(7))

doc.add_heading("10.2 Action Specification Cards (Detailed)", level=2)

add_body(doc,
    "For each action, expand the table row into a detailed specification card below. "
    "Use the card format for each action \u2014 copy and fill in the template for each "
    "action in your flow."
)

# Template card for a processing action
doc.add_heading("Action: [action_name]", level=3)

card_fields = [
    ("Purpose:", "One sentence stating the result this action delivers (derived from Output column of consolidated scope steps)."),
    ("Action Type:", "Data Gathering / LLM Analysis / LLM Generation / HIL Checkpoint"),
    ("Tools:", "List the APIs, LLM calls, or utility functions this action uses (traced from Data Inventory)."),
    ("Logic:", "Step-by-step description of what the action does. Derive from Decision Logic and Output columns. Be specific enough for a developer to implement, but do not write pseudocode."),
    ("Prompt Pattern (LLM actions only):", "What context the prompt receives (name specific memory fields and config data) and what output format it produces (JSON schema with named keys, Markdown sections, plain text). Do not write the actual prompt text."),
    ("Parallelism (data-gathering actions):", "Are the action's tool calls independent? Can they run concurrently?"),
    ("Surfaced Data (HIL actions):", "What the action presents to the human at the HIL checkpoint \u2014 which memory fields and in what format."),
    ("Resume Format (HIL actions):", "What the human provides when resuming the flow and how the action processes it."),
    ("Error Handling:", "What failures this action can encounter, how it detects them, and which response strategy it applies."),
    ("Input Memory Fields:", "Which memory fields this action reads."),
    ("Output Memory Fields:", "Which memory fields this action writes."),
    ("Edge Type:", "Unconditional \u2192 [next_action] or Conditional \u2192 {condition: target_action, ...}"),
]

for label, instruction in card_fields:
    p = doc.add_paragraph()
    run = p.add_run(label + " ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(instruction)
    run2.italic = True
    run2.font.size = Pt(10)

add_body(doc, "")  # spacer

add_italic_instruction(doc,
    "Copy the card template above for each action in your flow. Not all fields apply to "
    "every action type \u2014 Parallelism is only for data-gathering actions; Surfaced Data and "
    "Resume Format are only for HIL actions; Prompt Pattern is only for LLM actions. Leave "
    "inapplicable fields out rather than writing 'N/A'."
)

add_tip(doc,
    "The prompt pattern is the contract between design and build. It specifies what "
    "data the prompt receives and what structure the output must follow \u2014 without writing "
    "the actual prompt text. If the prompt pattern says 'analyse the data and return "
    "results', the developer must reverse-engineer the intent. Be specific: name memory "
    "fields, config sections, and output keys."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 11. HUMAN-IN-THE-LOOP INTERACTION DESIGN
# ════════════════════════════════════════════════════════════════════

doc.add_heading("11. Human-in-the-Loop Interaction Design", level=1)

doc.add_heading("11.1 HIL Pattern Choice Rationale", level=2)

add_body(doc,
    "Every HIL checkpoint falls into one of two patterns. The choice affects your flow "
    "structure (unconditional vs. conditional edges), memory schema (single optional "
    "feedback field vs. decision field + instructions field), and downstream action prompts."
)

# Pattern comparison table
pattern_table = doc.add_table(rows=1, cols=3)
pattern_table.style = "Table Grid"
add_header_row(pattern_table, ["Criterion", "Feedback-as-Context", "Feedback-as-Control-Flow"], Pt(9))

pattern_rows = [
    ("Purpose of checkpoint", "Enrich the next action's reasoning", "Gate whether output proceeds, gets revised, or gets reworked"),
    ("Does human response affect routing?", "No \u2014 unconditional edge forward", "Yes \u2014 conditional edges to different actions"),
    ("Memory field shape", "Single optional feedback string", "Decision field (constrained values) + optional instructions string"),
    ("Edge type", "Unconditional edge", "Conditional edge (routing function that inspects the decision field)"),
    ("Key question", "Does the human's response change which action runs next, or only what the next action knows?", "Same question \u2014 if routing changes, it's control-flow"),
]

for vals in pattern_rows:
    row = pattern_table.add_row()
    for i, v in enumerate(vals):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(v)
        run.font.size = Pt(8)

add_italic_instruction(doc,
    "For each HIL checkpoint in your flow, state which pattern it uses and why. Reference "
    "the key question: 'Does the human's response change which action runs next (control-flow) "
    "or only what the next action knows (context)?' Example: 'hil_review_analysis uses "
    "Feedback-as-Context because the flow always proceeds to synthesise regardless of the "
    "human's input \u2014 corrections enrich the synthesis prompt but do not change routing.'"
)

doc.add_heading("11.2 HIL Interaction Design Table", level=2)

add_body(doc,
    "Document the detailed interaction design for each HIL checkpoint action. One row per "
    "HIL action."
)

# HIL table — columns match the CSV
hil_headers = [
    "HIL Action\nName",
    "HIL\nPattern",
    "What Is Surfaced\nto Human",
    "Expected\nHuman Input",
    "Memory Fields\nWritten",
    "Routing After\nResume",
    "Decision\nOptions",
    "Design\nRationale",
]

table = doc.add_table(rows=1, cols=8)
table.style = "Table Grid"
add_header_row(table, hil_headers, Pt(7))

# Example row — feedback-as-context
ex_row = table.add_row()
add_example_row(table, 1, [
    "hil_review_analysis",
    "Feedback-as-Context",
    "Formatted summary of all three analyses: ticket (score/trend/key issues/SLA), "
    "usage (score/trend/adoption), sentiment (score/summary/key threads). "
    "Formatted for review readability, not raw JSON.",
    "Freeform text \u2014 confirm, correct scores, add missing context, flag data quality issues. "
    "No structured format required.",
    "analysis_human_feedback",
    "Unconditional \u2192 synthesise. Flow always moves forward.",
    "N/A \u2014 freeform input, no structured decision.",
    "Reviewer evaluates related dimensions together. Feedback-as-Context because corrections "
    "enrich synthesis without changing routing. Enables Consolidation Principle 2 downstream.",
], Pt(6))

# Example row — feedback-as-control-flow
ex_row2 = table.add_row()
add_example_row(table, 2, [
    "hil_review_report",
    "Feedback-as-Control-Flow",
    "Complete report package: full report_markdown (7 sections), exec_summary, "
    "health_score, risks, opportunities. Explicit decision options shown.",
    "Structured response: (1) decision \u2014 approved/revise/rework. "
    "(2) instructions \u2014 required for revise and rework, not for approved.",
    "report_review_decision (always); report_human_feedback (revise); rework_instructions (rework)",
    "Conditional: approved \u2192 END; revise \u2192 generate_report; rework \u2192 analyse_health",
    "Approved: report ready for distribution. Revise: regenerate report only. "
    "Rework: re-run analysis and everything downstream.",
    "Quality gate with three decision paths giving proportional control. "
    "Routing function validates decision value to prevent silent misrouting.",
], Pt(6))

add_empty_rows(table, 4, 8, Pt(7))

add_tip(doc,
    "If a checkpoint has only one possible next step regardless of what the human says, "
    "it is Feedback-as-Context. If the human can send the flow to different actions, it is "
    "Feedback-as-Control-Flow. Making this decision during design is much easier than "
    "retrofitting routing fields and conditional edges mid-build."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 12. ERROR HANDLING DESIGN
# ════════════════════════════════════════════════════════════════════

doc.add_heading("12. Error Handling Design", level=1)

doc.add_heading("12.1 Error Handling Strategy Rationale", level=2)

add_body(doc,
    "Walk each action in your flow and identify what can go wrong. Failure modes fall "
    "into three categories: external service failures (HTTP errors, timeouts, rate "
    "limits), data quality failures (empty results, unexpected formats, stale data), "
    "and LLM output failures (malformed JSON, refusals, hallucination). For each "
    "failure, choose one of four strategies: error marker and continue, retry with "
    "modification, escalate via an HIL checkpoint, or abort."
)

# Strategy reference table
strat_table = doc.add_table(rows=1, cols=3)
strat_table.style = "Table Grid"
add_header_row(strat_table, ["Strategy", "When to Use", "Example"], Pt(9))

strat_rows = [
    ("Error marker and continue", "Failure is recoverable; downstream actions can work with partial data", "Data source API down \u2014 store error marker, continue with available sources"),
    ("Retry with modification", "Failure is transient or fixable by changing the request; set a retry limit", "LLM returns malformed JSON \u2014 retry with stricter prompt (up to 2\u20133 attempts)"),
    ("Escalate via HIL checkpoint", "Automated recovery has failed or situation requires human judgement", "All retries exhausted on a critical action \u2014 escalate to an HIL checkpoint with failure context"),
    ("Mitigated by design", "Risk is addressed through prompt design, data grounding, or HIL verification rather than runtime detection", "LLM hallucination risk \u2014 prompts require evidence citations; HIL checkpoints catch unsupported claims"),
    ("Abort", "Continuing would produce garbage; failure is unrecoverable", "Authentication permanently revoked for the only data source"),
]

for vals in strat_rows:
    row = strat_table.add_row()
    for i, v in enumerate(vals):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(v)
        run.font.size = Pt(8)

add_italic_instruction(doc,
    "Describe your overall error handling philosophy for this workflow. Do you default "
    "to graceful degradation (partial results with flagged gaps)? Which actions are critical "
    "enough to warrant escalation if they fail? Are there any actions where abort is the "
    "correct response? How does your HIL checkpoint design serve as a safety net for "
    "errors that automated strategies cannot catch?"
)

add_warning(doc,
    "Default to graceful degradation. An agent that produces a partial result with flagged "
    "gaps is more useful than one that crashes. Reserve abort for cases where continuing "
    "would actively mislead the user. Every retry needs a maximum attempt count \u2014 without "
    "one, a persistent failure becomes an infinite loop."
)

doc.add_heading("12.2 Error Handling Table", level=2)

add_body(doc,
    "Document every identified failure mode. One row per failure mode. The five example "
    "rows below demonstrate all three failure categories (external service, data quality, "
    "LLM output) and all four response strategies."
)

# Error handling table — columns match the CSV
error_headers = [
    "Failure Mode",
    "Failure\nCategory",
    "Affected\nActions",
    "Detection\nMethod",
    "Response\nStrategy",
    "Strategy Detail",
    "Downstream\nImpact",
    "Notes",
]

table = doc.add_table(rows=1, cols=8)
table.style = "Table Grid"
add_header_row(table, error_headers, Pt(7))

# Example row 1 — Data source API failure (external service)
ex_row = table.add_row()
add_example_row(table, 1, [
    "Data source API failure (HTTP error / timeout / rate limit)",
    "External service",
    "gather_data",
    "HTTP error code (4xx/5xx) or timeout exception from API wrapper",
    "Error marker and continue",
    "Store error marker dict in memory field (e.g. {error: true, source: 'Support API', "
    "reason: 'HTTP 503'}). Continue with remaining sources. Do not retry \u2014 API failures "
    "are typically not transient within the same execution.",
    "analyse_health skips that dimension (score: 'Unknown'). synthesise flags gaps. "
    "generate_report includes data gap notice. hil_review_report: human sees gap.",
    "Default to graceful degradation. A report with one flagged gap is more useful than no report.",
], Pt(6))

# Example row 2 — Data source returns empty (data quality)
ex_row2 = table.add_row()
add_example_row(table, 2, [
    "Data source returns empty result set",
    "Data quality",
    "gather_data",
    "Valid API response with zero records \u2014 empty list [] or empty dict {}. "
    "Detected by checking response length after successful API call.",
    "Error marker and continue",
    "Store 'no data available' marker (e.g. {empty: true, source: 'Support API', "
    "note: 'No tickets found for account in this quarter'}). Distinct from an error marker: "
    "the API succeeded but returned no data. Do not infer meaning from absence.",
    "analyse_health notes 'no data available' for that dimension \u2014 does not attempt analysis. "
    "Report states 'No data available for this period' rather than omitting the section. "
    "hil_review_analysis: human sees gap and can add context.",
    "Absence of data is not evidence of absence of problems. The source may be misconfigured "
    "or the account may use a different channel. Flag gaps \u2014 don't guess.",
], Pt(6))

# Example row 3 — LLM malformed JSON (LLM output)
ex_row3 = table.add_row()
add_example_row(table, 3, [
    "LLM produces malformed JSON output",
    "LLM output",
    "analyse_health; synthesise",
    "JSON parse failure (json.JSONDecodeError) or schema validation failure (missing keys / wrong types)",
    "Retry with modification",
    "Retry with stricter format instruction (up to 2 retries per call, 3 total attempts). "
    "In analyse_health: each dimension retries independently. In synthesise: entire synthesis retries. "
    "If all retries exhausted: escalate via an HIL checkpoint with failure context.",
    "If escalated: human provides the structured data manually. Downstream actions operate "
    "normally on human-provided data.",
    "Retry budget (2 retries) prevents infinite loops. Stricter prompts on retry often resolve formatting issues.",
], Pt(6))

# Example row 4 — LLM hallucination risk (mitigated by design)
ex_row4 = table.add_row()
add_example_row(table, 4, [
    "LLM hallucination / unsupported claims in generated content",
    "LLM output",
    "analyse_health; synthesise; generate_report",
    "Not detectable at runtime via automated checks \u2014 mitigated by design rather than "
    "runtime detection.",
    "Mitigated by design (prompt grounding + HIL verification)",
    "All LLM prompts instruct the model to base conclusions only on provided data and cite "
    "evidence. Scoring criteria externalised to config. Two HIL checkpoints serve as human "
    "verification layers where domain experts catch unsupported claims.",
    "hil_review_analysis: reviewer compares analyses against raw data. hil_review_report: "
    "reviewer evaluates final report for unsupported claims. Revise path fixes report issues; "
    "rework path re-runs analysis.",
    "Not every failure mode needs a runtime strategy. Layered defence: prompt grounding reduces "
    "frequency, HIL checkpoints catch what slips through.",
], Pt(6))

# Example row 5 — Human doesn't respond (external service)
ex_row5 = table.add_row()
add_example_row(table, 5, [
    "Human does not respond at HIL checkpoint",
    "External service",
    "hil_review_analysis; hil_review_report",
    "Configurable timeout or indefinite wait; detected by application layer",
    "Error marker and continue (memory persistence)",
    "Memory persists via the chosen framework's checkpointing layer. Flow pauses indefinitely. "
    "Can be resumed at any time. No data loss. Application layer handles reminders.",
    "No downstream impact until human responds. All memory preserved exactly.",
    "Checkpointing in the chosen framework makes unresponsive humans a non-issue for flow integrity. "
    "Operational concern (SLA) handled by application layer.",
], Pt(6))

add_empty_rows(table, 6, 8, Pt(7))

add_tip(doc,
    "Design error handling at the action boundary, not the tool boundary. The same tool "
    "failure might warrant different responses in different actions \u2014 a CRM timeout during "
    "initial data gathering might be recoverable (continue with other sources), but the "
    "same timeout during a critical validation step might require escalation."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 13. COMPLETION CHECKLIST
# ════════════════════════════════════════════════════════════════════

doc.add_heading("13. Completion Checklist", level=1)

add_body(doc,
    "Review your completed document against this checklist. Every item should be true "
    "before you move to Stage 5 (Build)."
)

checklist_items = [
    ("Scope-to-Action Mapping", [
        "Every AUTOMATE and HUMAN-IN-THE-LOOP scope step appears in the mapping table",
        "MANUAL steps are listed as 'Outside flow'",
        "Split-boundary steps are decomposed (AUTOMATE component and HIL component mapped separately)",
        "Every consolidation decision has a written rationale citing a specific rule or principle",
        "HIL checkpoints are consolidated \u2014 the number of HIL checkpoints is less than the number of HIL-tagged scope steps",
    ]),
    ("Flow Structure", [
        "There is a clear path from start to end for every possible execution",
        "Conditional edges are labelled with their routing condition",
        "The pattern selection rationale explains why this flow structure was chosen",
        "The diagram matches the mapping table \u2014 every action in the table appears in the diagram",
    ]),
    ("Memory Schema", [
        "Every field traces to a specific scope step Output, Data Inventory entry, or HIL pattern",
        "Field types are explicit (not generic dict for everything)",
        "Reducer annotations are present where multiple actions write to the same field",
        "HIL feedback fields match the pattern: single optional feedback string for context, decision + instructions for control-flow",
        "The pseudocode memory schema matches the table",
    ]),
    ("Action Specifications", [
        "Every action in the flow structure has a specification card",
        "Each action's purpose can be stated in one sentence",
        "Tools are traced from the Data Inventory (no tools appear that aren't in the inventory)",
        "Logic steps are specific enough for independent implementation",
        "Prompt patterns specify input memory fields and output format (not vague 'analyse the data')",
        "Error handling is specified per action, not generically",
    ]),
    ("HIL Interaction Design", [
        "Every HIL checkpoint action has an interaction design entry",
        "The pattern choice (context vs. control-flow) is justified with the key question",
        "Surfaced data specifies which memory fields and in what format",
        "Resume format specifies what the human provides and how it's processed",
        "Control-flow checkpoints have explicit decision options with descriptions",
    ]),
    ("Error Handling", [
        "Every action's failure modes are identified across all three categories (external, data quality, LLM)",
        "Every failure mode has a detection method",
        "Every failure mode has a deliberate response strategy (not default 'crash')",
        "Retry strategies have maximum attempt counts",
        "At least one escalation path exists for critical failures",
    ]),
]

for section, items in checklist_items:
    p = doc.add_paragraph()
    run = p.add_run(section)
    run.bold = True
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run("\u2610  " + item)
        run.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════
# 14. NOTES & OPEN QUESTIONS
# ════════════════════════════════════════════════════════════════════

doc.add_heading("14. Notes & Open Questions", level=1)

add_italic_instruction(doc,
    "Use this section to capture anything you are unsure about, questions that need "
    "answers before Stage 5 can begin, design trade-offs you want to revisit, or "
    "assumptions that need validation. Flag items that are blocking (must be resolved "
    "before build) vs. non-blocking (can be resolved during build)."
)

# Blank lines for notes
for _ in range(8):
    doc.add_paragraph("")

# ════════════════════════════════════════════════════════════════════
# 15. NEXT STEPS
# ════════════════════════════════════════════════════════════════════

doc.add_heading("15. Next Steps", level=1)

add_body(doc,
    "Once this Design Document is complete and all checklist items are verified:"
)

steps = [
    ("Review with a developer (if you are not one).", "Walk through the flow structure, memory schema, and action specifications together. The developer should confirm that each action spec is implementable as written and that the memory schema is complete."),
    ("Resolve any open questions.", "Address blocking items in the Notes & Open Questions section before proceeding. Non-blocking items can be resolved during Stage 5."),
    ("Proceed to Stage 5: Build.", "The build stage translates this design into a working agent on the platform your team has chosen. The flow structure becomes the platform\u2019s flow construct. The memory schema becomes its typed memory/state object. Each action spec becomes a concrete action implementation. The HIL designs become pause-and-resume implementations."),
    ("Keep this document as a living reference.", "During build, you will discover edge cases and make implementation decisions not covered here. Update this document to reflect those decisions \u2014 it serves as the architectural record for Stage 6 (Evaluate & Iterate) when you trace evaluation findings back to design choices."),
]

for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {title} ")
    run.bold = True
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

add_tip(doc,
    "The quality of Stage 5 (Build) is entirely determined by the quality of this "
    "Design Document. The most common build failures trace back to vague action specs, "
    "incomplete memory schemas, or missing error handling strategies. Invest the time "
    "here \u2014 it pays off dramatically during implementation."
)


# ── Save ────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "Stage 4 - Design - Template.docx")
doc.save(output_path)
print(f"Created: {output_path}")
